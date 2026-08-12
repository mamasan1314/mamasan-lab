from __future__ import annotations

import sys
from datetime import datetime
from pathlib import Path


HERE = Path(__file__).resolve().parent
VENDOR = HERE / "vendor"
if VENDOR.exists():
    sys.path.insert(0, str(VENDOR))

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


PROJECT = HERE.parent
OUTPUTS = PROJECT / "outputs"
OUTPUTS.mkdir(parents=True, exist_ok=True)
DOCX = OUTPUTS / "希望之光_老師下次會議簡單討論稿_202608.docx"
LOGO = OUTPUTS / "Hope_Light_Logo_初稿_01.png"

NAVY = "17324D"
TEAL = "2F7D73"
MINT = "DCEFE9"
GOLD = "D8A84E"
PALE_GOLD = "F8EDD2"
PALE_RED = "F7DEDE"
LIGHT = "F5F7F8"
WHITE = "FFFFFF"
GRAY = "66717D"


POSTS = [
    {
        "number": 1,
        "title": "歡迎來到 Hope Light Moment｜這裡不只是賣蠟燭",
        "format": "Reels｜帳號開場",
        "caption": (
            "歡迎來到 Hope Light Moment。這裡記錄的，不是一顆蠟燭有多神奇，而是你願意為自己停下來的那五分鐘。"
            "我們會分享頻率儀式蠟燭、大小貴人的生活用法、老師的日常，以及一個個真實的希望時刻。"
            "願這道光，陪你在忙亂中重新看見自己，找回今天的節奏。"
        ),
        "shots": "Logo出現 → 老師點燭 → 寫下「今天我想回到＿＿狀態」",
        "cta": "留言告訴我：工作、關係、學習、睡前，你現在最想整理哪個時刻？",
        "note": "建議作為第1篇置頂貼文。",
    },
    {
        "number": 2,
        "title": "我們的蠟燭，和別人不一樣的地方",
        "format": "Reels｜製作幕後",
        "caption": (
            "一顆蠟燭製作完成，不代表它已經準備好出發。對 Hope Light 來說，重要的是蠟燭回到老師手上後，"
            "還會經過確認、整理、淨化與頻率設定。蠟燭只是媒介，真正想交到你手上的，是一個願意停下來、"
            "重新整理自己的提醒。這就是一顆 Hope Light 蠟燭出生的故事。"
        ),
        "shots": "成品回到工作桌 → 老師整理／設定的手部畫面 → 包裝完成",
        "cta": "追蹤下一篇，一起看一顆蠟燭如何準備出發。",
        "note": "發布前由老師確認實際流程與「淨化／頻率設定」可公開畫面。",
    },
    {
        "number": 3,
        "title": "頻率，為什麼會影響一個人的每一天？",
        "format": "輪播｜生活觀點",
        "caption": (
            "我們說的頻率，不是按下一個按鈕就改變命運，而是你每天反覆停留的狀態。焦急時，我們容易只看到問題；"
            "比較穩定時，才有空間看見選擇。點燃蠟燭不是把答案交給火焰，而是用一個固定動作提醒自己：先停一下、"
            "呼吸、看見現在，再決定下一步。蠟燭是媒介，你的覺察與行動才是核心。"
        ),
        "shots": "第1張提問 → 第2張焦急／穩定對照 → 第3張五分鐘儀式 → 第4張下一步",
        "cta": "收藏這篇。今天你想讓自己回到什麼狀態？",
        "note": "不使用腦波、赫茲或生理改變等未經驗證說法。",
    },
    {
        "number": 4,
        "title": "點燃一天，不是多做一件事，而是先選擇今天怎麼開始",
        "format": "Reels｜五分鐘教學",
        "caption": (
            "儀式感不需要很複雜。把手機放遠、整理一小塊桌面、點燃蠟燭，慢慢呼吸三次，再寫下今天最重要的一件事。"
            "五分鐘後，你不一定立刻變得完美，但你已經替今天做了一個清楚的開始。每天重複的小動作，"
            "會慢慢變成陪伴自己的方式。使用明火時請保持通風，遠離易燃物、孩子與寵物。"
        ),
        "shots": "手機翻面 → 點燭與呼吸 → 寫下今日一件事",
        "cta": "收藏這個五分鐘流程，明天早上一起試一次。",
        "note": "畫面不可把點燃中的蠟燭留在無人空間。",
    },
    {
        "number": 5,
        "title": "我為什麼開始做蠟燭？最初只是想陪孩子好好開始學習",
        "format": "Reels｜老師故事",
        "caption": (
            "一開始，我不是想做一個商品。那時看著孩子準備學習，我心裡只有一個很單純的願望：希望他能有一個比較安定、"
            "更有效率的開始。於是我開始做蠟燭，也慢慢發現，真正有價值的不只是蠟燭，而是我們一起整理桌面、說清楚今天要完成什麼、"
            "陪彼此進入狀態的那段時間。蠟燭不會讓成績自動變好，但可以成為學習開始前的固定提醒。"
        ),
        "shots": "老師說故事 → 孩子學習空間／無人物空鏡 → 點燭後寫下學習目標",
        "cta": "你想為孩子建立一個怎樣的學習開場？",
        "note": "老師須確認故事細節；若出現孩子或個資，先取得同意。",
    },
    {
        "number": 6,
        "title": "1號 防護除穢｜替空間與自己重新畫出界線",
        "format": "輪播｜10款蠟燭 1/10",
        "caption": (
            "當外界訊息太多、空間很亂、心裡一直被打擾，1號蠟燭可以作為整理環境與界線的開始。點燃後，先清出一小塊空間，"
            "關掉不必要的通知，寫下今天不再讓什麼持續消耗自己。所謂防護，不是控制外界，而是提醒自己把注意力與選擇權帶回來。"
        ),
        "shots": "雜亂桌面 → 點燃1號 → 整理後寫下界線",
        "cta": "今天你最想清掉的是雜物、雜訊，還是一件反覆消耗你的事？",
        "note": "對外以環境整理與界線表達，不保證驅除特定人事物。",
    },
    {
        "number": 7,
        "title": "2號 財富豐盛｜把焦慮換成一個清楚的金錢行動",
        "format": "Reels｜10款蠟燭 2/10",
        "caption": (
            "財富不是只靠等待。點燃2號蠟燭時，問自己三件事：我現在最擔心什麼？手上有哪些資源？今天能完成哪一個實際行動？"
            "可能是整理帳目、追蹤一位客戶，或停止一筆不必要的花費。這顆蠟燭的角色，是提醒自己從焦慮回到清楚與行動。"
        ),
        "shots": "帳目／筆記 → 點燃2號 → 圈出今日金錢行動",
        "cta": "留言「行動」，寫下你今天願意完成的一件事。",
        "note": "不保證招財、獲利或投資結果。",
    },
    {
        "number": 8,
        "title": "3號 吸引顧客｜先想清楚，今天要幫助誰？",
        "format": "輪播｜10款蠟燭 3/10",
        "caption": (
            "當生意卡住時，我們很容易只想著「怎麼讓更多人買」。點燃3號蠟燭前，先換一個問題：我今天最想幫助哪一種人？"
            "他正在困擾什麼？我能把哪一件事說得更清楚？接著完成一個顧客行動：回覆訊息、整理介紹，或真心關心一位舊客。"
        ),
        "shots": "顧客問題便條 → 點燃3號 → 回覆一則訊息",
        "cta": "收藏這三個問題，下一次發文前先回答一次。",
        "note": "不保證帶來顧客或成交。",
    },
    {
        "number": 9,
        "title": "4號 吸引桃花｜先讓自己回到願意連結的狀態",
        "format": "Reels｜10款蠟燭 4/10",
        "caption": (
            "想遇見好的關係之前，也可以先問：我有沒有好好照顧自己？我願不願意讓別人看見真實的我？點燃4號蠟燭，"
            "留五分鐘寫下自己期待的關係，以及自己也願意付出的相處方式。它不是保證某個人出現，而是提醒自己帶著清楚與開放走進關係。"
        ),
        "shots": "整理自己／空間 → 點燃4號 → 寫下理想關係的三個詞",
        "cta": "你希望一段關係帶給彼此什麼感受？",
        "note": "不保證桃花、復合或特定關係結果。",
    },
    {
        "number": 10,
        "title": "5號 貴人常臨｜看見支持，也練習開口",
        "format": "輪播｜10款蠟燭 5/10",
        "caption": (
            "有時候不是身邊沒有貴人，而是我們太習慣一個人撐著。點燃5號蠟燭時，想起三位曾經幫助你的人，"
            "對其中一位說聲謝謝；再寫下一件你願意開口請教的事。貴人的意義，不只是等人來拉你一把，也包括你願意連結、回應與合作。"
        ),
        "shots": "寫下三個名字 → 點燃5號 → 傳出一則感謝訊息",
        "cta": "今天，向一位曾幫助你的人說謝謝。",
        "note": "不保證特定人物或機會出現。",
    },
    {
        "number": 11,
        "title": "6號 學業進步｜替學習建立一個固定的開始",
        "format": "Reels｜10款蠟燭 6/10",
        "caption": (
            "讀書前最難的，常常不是不會，而是還沒有開始。點燃6號蠟燭後，先把桌面留下這次需要的東西，"
            "寫下一個小目標，再設定一段專心時間。固定的開始動作，可以幫助孩子或大人知道：現在要進入學習了。"
            "它不保證成績，而是陪你建立能重複的學習節奏。"
        ),
        "shots": "清桌面 → 點燃6號 → 寫目標並開始計時",
        "cta": "收藏起來，今晚和孩子一起試一次三步學習開場。",
        "note": "不宣稱改善ADHD、專注力或保證成績。",
    },
    {
        "number": 12,
        "title": "7號 小人退散｜不控制別人，先保護自己的界線",
        "format": "輪播｜10款蠟燭 7/10",
        "caption": (
            "面對讓人不舒服的關係，我們未必能改變對方，但可以決定自己怎麼回應。點燃7號蠟燭，寫下三件事："
            "我不再接受什麼？我要保持什麼距離？遇到同樣情況時，我準備怎麼說？把「退散」翻成生活語言，"
            "就是停止反覆消耗，練習清楚而平靜的界線。"
        ),
        "shots": "關掉干擾訊息 → 點燃7號 → 寫下一句界線說法",
        "cta": "寫下一句你想練習的界線，不必公開內容。",
        "note": "不影射或攻擊任何特定人物。",
    },
    {
        "number": 13,
        "title": "8號 好運爆棚｜好運來時，你有沒有準備好看見？",
        "format": "Reels｜10款蠟燭 8/10",
        "caption": (
            "好運不一定是突然中獎，也可能是一封訊息、一個邀請，或你終於願意開始的念頭。點燃8號蠟燭，"
            "寫下今天已經出現的三個小機會，再選一個能立刻回應的行動。這顆蠟燭想提醒的是：保持開放，也準備好用行動接住機會。"
        ),
        "shots": "一天中的小機會 → 點燃8號 → 完成一個回應行動",
        "cta": "今天出現過哪一個你差點忽略的小機會？",
        "note": "不保證改運、中獎或特定結果。",
    },
    {
        "number": 14,
        "title": "9號 感情升溫｜讓五分鐘重新有對話",
        "format": "輪播｜10款蠟燭 9/10",
        "caption": (
            "關係變淡，很多時候不是沒有感情，而是彼此太久沒有真正停下來。點燃9號蠟燭，把手機放遠，"
            "輪流說一件今天的感受，再說一句對對方的感謝。五分鐘不一定能解決所有問題，但可以成為重新開始聽見彼此的入口。"
        ),
        "shots": "兩杯水／兩個座位 → 點燃9號 → 寫下一句感謝",
        "cta": "今晚試著問對方：今天有哪一刻，你希望我懂？",
        "note": "不保證復合或改善所有關係；涉及衝突與安全時應尋求適當協助。",
    },
    {
        "number": 15,
        "title": "10號 舒壓好眠｜替今天做一個溫柔的收尾",
        "format": "Reels｜10款蠟燭 10/10",
        "caption": (
            "睡前腦中還有很多事情時，先不要逼自己立刻放空。把燈光調暗、點燃10號蠟燭，慢慢呼吸，"
            "把還沒完成的事寫到明天，再記下一件今天已經做好的事。這是一個讓身體知道「今天要收尾了」的生活儀式。"
            "蠟燭不治療失眠；睡前務必完全熄滅，不能在睡著或無人看管時燃燒。"
        ),
        "shots": "關閉螢幕 → 點燃10號並寫明日清單 → 完全熄滅蠟燭",
        "cta": "收藏這個晚間流程，今晚替自己留五分鐘。",
        "note": "不宣稱治療失眠、焦慮或其他身心狀況。",
    },
]


def set_cell_shading(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, top: int = 100, start: int = 130, bottom: int = 100, end: int = 130) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def add_page_number(paragraph) -> None:
    paragraph.add_run("第 ")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)
    paragraph.add_run(" 頁")


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.25)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.35)
    section.right_margin = Cm(1.35)

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft JhengHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    styles["Normal"].font.size = Pt(10)
    styles["Normal"].paragraph_format.space_after = Pt(4)
    for style_name, size, color in (
        ("Title", 24, NAVY),
        ("Subtitle", 12, TEAL),
        ("Heading 1", 17, NAVY),
        ("Heading 2", 13, TEAL),
    ):
        style = styles[style_name]
        style.font.name = "Microsoft JhengHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = style_name != "Subtitle"
        style.paragraph_format.space_before = Pt(5)
        style.paragraph_format.space_after = Pt(5)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("HOPE LIGHT｜下次會議討論稿")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(GRAY)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("希望之光 Hope Light｜[Codex]｜")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(GRAY)
    add_page_number(footer)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        set_cell_shading(cell, TEAL)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(8.5)
    set_repeat_table_header(table.rows[0])
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for col_index, value in enumerate(row):
            cells[col_index].text = str(value)
            set_cell_shading(cells[col_index], WHITE if row_index % 2 == 0 else LIGHT)
            set_cell_margins(cells[col_index])
            cells[col_index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cells[col_index].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Cm(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.add_run(item)


def add_box(doc: Document, text: str, color: str = MINT, bold_prefix: str | None = None) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, color)
    set_cell_margins(cell, 150, 180, 150, 180)
    paragraph = cell.paragraphs[0]
    if bold_prefix and text.startswith(bold_prefix):
        paragraph.add_run(bold_prefix).bold = True
        paragraph.add_run(text[len(bold_prefix) :])
    else:
        paragraph.add_run(text)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_post_card(doc: Document, post: dict[str, object]) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_shading(cell, WHITE)
    set_cell_margins(cell, 105, 150, 105, 150)
    cell.width = Cm(18.1)

    title = cell.paragraphs[0]
    title.paragraph_format.space_after = Pt(2)
    title.paragraph_format.keep_with_next = True
    run = title.add_run(f'{post["number"]}. {post["title"]}')
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(NAVY)

    meta = cell.add_paragraph()
    meta.paragraph_format.space_after = Pt(2)
    run = meta.add_run(str(post["format"]))
    run.bold = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(TEAL)

    caption = cell.add_paragraph()
    caption.paragraph_format.space_after = Pt(2)
    caption.add_run("文案｜").bold = True
    caption.add_run(str(post["caption"]))
    for run in caption.runs:
        run.font.size = Pt(8.6)

    shooting = cell.add_paragraph()
    shooting.paragraph_format.space_after = Pt(1)
    shooting.add_run("簡單拍法｜").bold = True
    shooting.add_run(str(post["shots"]))
    for run in shooting.runs:
        run.font.size = Pt(8.2)

    cta = cell.add_paragraph()
    cta.paragraph_format.space_after = Pt(1)
    cta.add_run("CTA｜").bold = True
    cta.add_run(str(post["cta"]))
    for run in cta.runs:
        run.font.size = Pt(8.2)

    note = cell.add_paragraph()
    note.paragraph_format.space_after = Pt(0)
    note.add_run("發布前確認｜").bold = True
    note.add_run(str(post["note"]))
    for run in note.runs:
        run.font.size = Pt(7.8)
        run.font.color.rgb = RGBColor.from_string(GRAY)

    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def build_document() -> None:
    if not LOGO.exists():
        raise FileNotFoundError(f"Logo not found: {LOGO}")

    doc = Document()
    configure_document(doc)
    props = doc.core_properties
    props.title = "希望之光下次會議討論稿｜IG與15篇內容"
    props.subject = "一人工作室下次會議、Hope Light Moment與首批15篇內容"
    props.author = "[Codex]"
    props.created = datetime(2026, 8, 9)

    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover.add_run().add_picture(str(LOGO), width=Cm(8.2))
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("下次會議討論稿\nIG 帳號與首批15篇內容")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Hope Light Moment｜一人工作室簡單版｜2026-08-09｜[Codex]")
    add_box(
        doc,
        "這次會議不談複雜系統，只要定下帳號門面、Logo方向、蠟燭故事與第一批內容。會後老師只需要集中拍一次、確認文字；其餘由品牌規劃者整理。",
        MINT,
    )
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("會議目標：60分鐘內做完5個決定，讓帳號可以開始。")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(NAVY)
    doc.add_page_break()

    doc.add_heading("一、下次會議只討論這5件事", level=1)
    add_table(
        doc,
        ["時間", "討論內容", "會議結束要有的答案"],
        [
            ["10分鐘", "Hope Light Moment帳號門面", "帳號名稱、使用者名稱、Bio、LINE連結"],
            ["10分鐘", "Logo初稿", "保留方向或提出1次明確修改"],
            ["15分鐘", "蠟燭差異與老師故事", "確認製作／淨化／頻率設定流程，以及孩子學習故事"],
            ["15分鐘", "15篇內容", "確認哪些可以直接發、哪些要補照片或故事"],
            ["10分鐘", "拍攝與分工", "確定60分鐘拍攝日期、老師要準備的物品與核稿方式"],
        ],
        [2.5, 6.0, 9.5],
    )

    doc.add_heading("1.1 老師會前只要準備", level=2)
    add_bullets(
        doc,
        [
            "10款蠟燭各1顆，或每款一張清楚照片。",
            "老師實際處理蠟燭的流程：誰製作、回來後做什麼、哪些步驟可以拍。",
            "「為了孩子學習而開始做蠟燭」的真實時間、情境與可公開程度。",
            "1至3個可匿名分享的客戶故事；公開前先確認當事人同意。",
            "官方LINE或預約連結，以及希望客戶私訊時使用的關鍵字。",
        ],
    )

    doc.add_heading("1.2 會後誰做什麼", level=2)
    add_table(
        doc,
        ["角色", "下一步", "完成標準"],
        [
            ["老師", "集中拍攝60分鐘、補流程與故事、一次核准文字", "每週不超過60-90分鐘"],
            ["品牌規劃者", "整理帳號、剪輯、封面、排程與私訊引導", "每週固定2篇，逐步發完15篇"],
        ],
        [3.0, 8.0, 7.0],
    )
    add_box(doc, "會議結束後第一個動作：先完成帳號門面，再發布3篇置頂內容，不必一次把15篇全部做好。", PALE_GOLD)
    doc.add_page_break()

    doc.add_heading("二、Hope Light Moment IG帳號規劃", level=1)
    add_table(
        doc,
        ["項目", "建議"],
        [
            ["帳號角色", "生活化主帳號：頻率蠟燭、大小貴人、五分鐘儀式、老師故事與使用者故事"],
            ["建議使用者名稱", "@hopelight.moment（建立帳號時再確認是否可用）"],
            ["顯示名稱", "Hope Light Moment｜希望時刻"],
            ["頭像", "使用本次 Hope Light Logo；正式啟用前再製作純圖示小尺寸版本"],
            ["每週頻率", "每週2篇：1支Reels＋1篇照片／輪播；忙碌時至少保留1篇"],
            ["內容順序", "先說品牌與故事，再介紹10款蠟燭；大小貴人放在下一階段"],
            ["唯一導流", "貼文只引導私訊一個關鍵字或加入LINE，不同時放多個選項"],
        ],
        [4.0, 14.0],
    )

    doc.add_heading("2.1 IG帳號介紹文（Bio）", level=2)
    add_box(
        doc,
        "Hope Light Moment｜希望時刻\n🕯️ 每天五分鐘，陪你找回自己的節奏\n頻率儀式蠟燭・大小貴人・真實故事\n👇 私訊「希望」／加入LINE找到適合你的儀式",
        MINT,
    )

    doc.add_heading("2.2 帳號打開後先做", level=2)
    add_table(
        doc,
        ["位置", "只放這些"],
        [
            ["置頂3篇", "①這個帳號在做什麼　②蠟燭不同在哪裡　③十款蠟燭怎麼選"],
            ["精選動態4個", "第一次來／10款蠟燭／五分鐘儀式／真實故事"],
            ["視覺配色", "深藍 #17324D、暖金 #D8A84E、暖白；畫面保持安靜、溫暖、留白"],
            ["封面字數", "每張只留8至14字；不用塞滿功效與說明"],
        ],
        [4.0, 14.0],
    )
    add_box(doc, "一人工作室原則：不每天發、不追所有熱門題目、不為兩個帳號做兩套內容。先把 Moment 做好。", PALE_GOLD)
    doc.add_page_break()

    doc.add_heading("三、Hope Light Logo 初稿", level=1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(LOGO), width=Cm(5.3))
    add_table(
        doc,
        ["元素", "代表意思"],
        [
            ["火焰與向上光線", "希望、點燃與每天重新開始"],
            ["開放圓形／地平線", "承接生活中的不同時刻，也保留前進空間"],
            ["深藍＋暖金", "深藍代表穩定與信任；暖金代表光與溫度"],
            ["簡潔英文名稱", "讓 Hope Light 可延伸到蠟燭、大小貴人與其他服務"],
        ],
        [4.5, 13.5],
    )
    add_box(
        doc,
        "會議只需要回答：①喜不喜歡這個符號方向？②顏色是否保留？③文字只留 HOPE LIGHT，還是要加中文「希望之光」？確認後再做純圖示頭像、透明底與黑白版。",
        PALE_GOLD,
    )

    doc.add_heading("3.1 15篇的簡單發布順序", level=2)
    add_table(
        doc,
        ["週次", "第1篇", "第2篇"],
        [
            ["第1週", "1 帳號開場", "2 蠟燭差異"],
            ["第2週", "3 頻率的重要", "4 五分鐘儀式"],
            ["第3週", "5 老師開始做蠟燭的原因", "6 1號 防護除穢"],
            ["第4週", "7 2號 財富豐盛", "8 3號 吸引顧客"],
            ["第5週", "9 4號 吸引桃花", "10 5號 貴人常臨"],
            ["第6週", "11 6號 學業進步", "12 7號 小人退散"],
            ["第7週", "13 8號 好運爆棚", "14 9號 感情升溫"],
            ["第8週", "15 10號 舒壓好眠", "休息／重發表現最好的一篇"],
        ],
        [2.5, 7.75, 7.75],
    )
    for start in range(0, len(POSTS), 3):
        end = min(start + 3, len(POSTS))
        doc.add_heading(f"四、15篇文章草稿｜第 {start + 1}–{end} 篇", level=1)
        if start == 0:
            add_box(
                doc,
                "使用方式：老師只確認事實與語氣；品牌規劃者再依照片長度微調。十款蠟燭雖保留產品名稱，但文章以生活時刻與可採取的行動表達，不把個案感受寫成保證功效。",
                MINT,
            )
        for post in POSTS[start:end]:
            add_post_card(doc, post)
        if end < len(POSTS):
            doc.add_page_break()

    doc.add_page_break()
    doc.add_heading("五、會議最後確認清單", level=1)
    add_table(
        doc,
        ["□", "老師要確認的答案"],
        [
            ["□", "帳號名稱、使用者名稱、Bio與LINE連結"],
            ["□", "Logo方向、顏色與是否加中文名稱"],
            ["□", "10款蠟燭正式名稱與每款可公開使用情境"],
            ["□", "蠟燭製作、淨化與頻率設定的真實流程"],
            ["□", "孩子學習故事與客戶故事的公開同意"],
            ["□", "60分鐘拍攝日期，以及第一週先發哪兩篇"],
        ],
        [1.2, 16.8],
    )
    add_box(
        doc,
        "發布安全提醒：不保證招財、桃花、顧客、成績、好運或睡眠結果；不宣稱改善疾病、腦波或生理數值。明火使用須遠離易燃物、孩子與寵物，睡前完全熄滅。",
        PALE_RED,
    )

    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    build_document()
