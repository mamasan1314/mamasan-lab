from __future__ import annotations

import sys
from datetime import date
from pathlib import Path


HERE = Path(__file__).resolve().parent
VENDOR = HERE / "vendor"
if VENDOR.exists():
    sys.path.insert(0, str(VENDOR))

import xlsxwriter
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from deliverable_data import PRODUCTS, PRODUCT_COPY


PROJECT = HERE.parent
OUTPUTS = PROJECT / "outputs"
OUTPUTS.mkdir(parents=True, exist_ok=True)
SIMPLE_XLSX = OUTPUTS / "希望之光_簡易產品報價庫存介紹表_202608.xlsx"
MOMENT_DOCX = OUTPUTS / "希望之光_一人工作室簡易執行表_202608.docx"
ROADMAP_DOCX = OUTPUTS / "希望之光_品牌規劃路徑圖_202608.docx"

NAVY = "17324D"
TEAL = "2F7D73"
MINT = "DCEFE9"
SKY = "EAF2F8"
GOLD = "D8A84E"
PALE_GOLD = "F8EDD2"
RED = "B54B4B"
PALE_RED = "F7DEDE"
GREEN = "3F7D52"
PALE_GREEN = "DDEEDC"
LIGHT = "F5F7F8"
WHITE = "FFFFFF"
GRAY = "66717D"


def price_status(sku: str, price: object) -> str:
    if sku == "HL-SV-CONS60":
        return "已確認，可報價"
    if sku in {"HL-DV-BEN-09", "HL-DV-SML-09", "HL-DV-BEN-07"}:
        return "價格衝突，暫勿對外"
    if price in (None, ""):
        return "價格未提供"
    return "來源售價，待老師核准"


def excel_formats(workbook: xlsxwriter.Workbook) -> dict[str, xlsxwriter.format.Format]:
    base = {"font_name": "Microsoft JhengHei", "font_size": 10, "valign": "vcenter"}
    return {
        "title": workbook.add_format(
            {**base, "bold": True, "font_size": 18, "font_color": WHITE, "bg_color": NAVY, "align": "left"}
        ),
        "note": workbook.add_format(
            {**base, "font_color": GRAY, "bg_color": SKY, "text_wrap": True, "align": "left"}
        ),
        "header": workbook.add_format(
            {**base, "bold": True, "font_color": WHITE, "bg_color": TEAL, "border": 1, "align": "center", "text_wrap": True}
        ),
        "text": workbook.add_format({**base, "border": 1, "text_wrap": True}),
        "center": workbook.add_format({**base, "border": 1, "align": "center", "text_wrap": True}),
        "input": workbook.add_format({**base, "border": 1, "bg_color": PALE_GOLD, "align": "center"}),
        "money": workbook.add_format({**base, "border": 1, "num_format": "#,##0;[Red]-#,##0"}),
        "money_input": workbook.add_format(
            {**base, "border": 1, "bg_color": PALE_GOLD, "num_format": "#,##0;[Red]-#,##0"}
        ),
        "date_input": workbook.add_format(
            {**base, "border": 1, "bg_color": PALE_GOLD, "num_format": "yyyy-mm-dd", "align": "center"}
        ),
    }


def setup_sheet(
    worksheet: xlsxwriter.worksheet.Worksheet,
    formats: dict[str, xlsxwriter.format.Format],
    title: str,
    note: str,
    headers: list[str],
    widths: list[float],
) -> None:
    worksheet.hide_gridlines(2)
    worksheet.set_tab_color(TEAL)
    worksheet.set_landscape()
    worksheet.fit_to_pages(1, 0)
    worksheet.set_margins(0.35, 0.35, 0.5, 0.5)
    worksheet.set_header(f'&L&"Microsoft JhengHei,Bold"{title}&R希望之光')
    worksheet.set_footer("&C第 &P / &N 頁")
    worksheet.merge_range(0, 0, 0, len(headers) - 1, title, formats["title"])
    worksheet.merge_range(1, 0, 1, len(headers) - 1, note, formats["note"])
    worksheet.set_row(0, 30)
    worksheet.set_row(1, 34)
    worksheet.set_row(3, 34)
    for col, (header, width) in enumerate(zip(headers, widths)):
        worksheet.write(3, col, header, formats["header"])
        worksheet.set_column(col, col, width)
    worksheet.freeze_panes(4, 0)
    worksheet.repeat_rows(0, 3)


def _build_detailed_simple_workbook() -> None:
    workbook = xlsxwriter.Workbook(SIMPLE_XLSX)
    workbook.set_properties(
        {
            "title": "希望之光簡易產品報價、庫存與介紹表",
            "subject": "品牌營運簡易工具",
            "author": "[Codex]",
            "company": "希望之光 Hope Light",
            "created": date(2026, 8, 9),
        }
    )
    formats = excel_formats(workbook)

    quote_headers = ["產品編號", "類別", "產品名稱", "單位", "單價", "價格狀態", "數量", "小計", "客戶／內部備註"]
    quote = workbook.add_worksheet("產品報價")
    setup_sheet(
        quote,
        formats,
        "希望之光｜簡易產品報價表",
        "黃底欄位可輸入數量。只有「已確認，可報價」可直接對外；其他價格先由老師核准。金額單位：新台幣。",
        quote_headers,
        [17, 13, 34, 8, 12, 22, 9, 14, 42],
    )
    quote_rows = []
    for product in PRODUCTS:
        sku, category, name = product[0], product[1], product[2]
        price, unit, note = product[6], product[9], product[14]
        quote_rows.append([sku, category, name, unit, price, price_status(sku, price), "", "", note])
    for offset, row in enumerate(quote_rows):
        excel_row = 5 + offset
        for col, value in enumerate(row):
            if col == 4:
                quote.write(4 + offset, col, value, formats["money"])
            elif col == 6:
                quote.write(4 + offset, col, value, formats["input"])
            elif col == 7:
                quote.write_formula(
                    4 + offset,
                    col,
                    f'=IF(OR(E{excel_row}="",G{excel_row}=""),"",E{excel_row}*G{excel_row})',
                    formats["money"],
                    "",
                )
            elif col in {0, 1, 3, 5}:
                quote.write(4 + offset, col, value, formats["center"])
            else:
                quote.write(4 + offset, col, value, formats["text"])
        quote.set_row(4 + offset, 42)
    quote.add_table(
        3,
        0,
        3 + len(quote_rows),
        len(quote_headers) - 1,
        {"name": "SimpleQuote", "style": "Table Style Medium 4", "columns": [{"header": h} for h in quote_headers]},
    )
    quote.data_validation(4, 6, 3 + len(quote_rows), 6, {"validate": "integer", "criteria": ">=", "value": 0})
    quote.conditional_format(4, 5, 3 + len(quote_rows), 5, {"type": "text", "criteria": "containing", "value": "已確認", "format": workbook.add_format({"bg_color": PALE_GREEN, "font_color": GREEN})})
    quote.conditional_format(4, 5, 3 + len(quote_rows), 5, {"type": "text", "criteria": "containing", "value": "衝突", "format": workbook.add_format({"bg_color": PALE_RED, "font_color": RED})})
    quote.conditional_format(4, 5, 3 + len(quote_rows), 5, {"type": "text", "criteria": "containing", "value": "待", "format": workbook.add_format({"bg_color": PALE_GOLD, "font_color": NAVY})})
    quote.print_area(0, 0, 3 + len(quote_rows), len(quote_headers) - 1)

    inventory_headers = [
        "產品編號",
        "類別",
        "產品名稱",
        "單位",
        "期初庫存",
        "本期入庫",
        "本期售出",
        "預留數量",
        "現有庫存",
        "可售庫存",
        "安全庫存",
        "補貨提醒",
        "盤點日期",
        "備註",
    ]
    inventory = workbook.add_worksheet("產品庫存")
    setup_sheet(
        inventory,
        formats,
        "希望之光｜簡易產品庫存表",
        "黃底欄位由實際盤點與交易填寫；本表沒有匯入舊庫存數字。現有庫存＝期初＋入庫－售出；可售庫存＝現有－預留。",
        inventory_headers,
        [17, 13, 32, 8, 11, 11, 11, 11, 11, 11, 11, 12, 13, 36],
    )
    inventory_products = [product for product in PRODUCTS if product[10] == "是"]
    for offset, product in enumerate(inventory_products):
        sheet_row = 4 + offset
        excel_row = sheet_row + 1
        sku, category, name, unit, safety, note = product[0], product[1], product[2], product[9], product[11], product[14]
        values = [sku, category, name, unit]
        for col, value in enumerate(values):
            inventory.write(sheet_row, col, value, formats["center"] if col in {0, 1, 3} else formats["text"])
        for col in range(4, 8):
            inventory.write_blank(sheet_row, col, None, formats["input"])
        inventory.write_formula(
            sheet_row,
            8,
            f'=IF(COUNT(E{excel_row}:G{excel_row})=0,"",SUM(E{excel_row}:F{excel_row})-G{excel_row})',
            formats["center"],
            "",
        )
        inventory.write_formula(sheet_row, 9, f'=IF(I{excel_row}="","",I{excel_row}-H{excel_row})', formats["center"], "")
        inventory.write(sheet_row, 10, safety, formats["input"])
        inventory.write_formula(
            sheet_row,
            11,
            f'=IF(J{excel_row}="","",IF(J{excel_row}<=K{excel_row},"需補貨","正常"))',
            formats["center"],
            "",
        )
        inventory.write_blank(sheet_row, 12, None, formats["date_input"])
        if sku == "HL-BD-21":
            note = f"{note}；套組建議另建立組成扣庫規則"
        inventory.write(sheet_row, 13, note, formats["text"])
        inventory.set_row(sheet_row, 42)
    inventory.add_table(
        3,
        0,
        3 + len(inventory_products),
        len(inventory_headers) - 1,
        {"name": "SimpleInventory", "style": "Table Style Medium 4", "columns": [{"header": h} for h in inventory_headers]},
    )
    inventory.data_validation(4, 4, 3 + len(inventory_products), 10, {"validate": "integer", "criteria": ">=", "value": 0})
    inventory.data_validation(4, 12, 3 + len(inventory_products), 12, {"validate": "date", "criteria": "between", "minimum": date(2020, 1, 1), "maximum": date(2100, 12, 31)})
    inventory.conditional_format(4, 11, 3 + len(inventory_products), 11, {"type": "text", "criteria": "containing", "value": "需補貨", "format": workbook.add_format({"bg_color": PALE_RED, "font_color": RED, "bold": True})})
    inventory.print_area(0, 0, 3 + len(inventory_products), len(inventory_headers) - 1)

    intro_headers = [
        "產品編號",
        "類別",
        "產品名稱",
        "一句話介紹",
        "適合情境",
        "建議使用方式",
        "注意事項",
        "售價",
        "建議CTA",
        "內容狀態",
        "待補資料",
    ]
    intro = workbook.add_worksheet("產品介紹")
    setup_sheet(
        intro,
        formats,
        "希望之光｜簡易產品介紹表",
        "本表以生活情境、整理與陪伴語言撰寫；避免治療、診斷、改運或保證效果。內容狀態未核准者不可直接複製上架。",
        intro_headers,
        [17, 13, 30, 44, 34, 44, 52, 12, 28, 16, 44],
    )
    product_map = {product[0]: product for product in PRODUCTS}
    for offset, copy in enumerate(PRODUCT_COPY):
        sku = copy[0]
        product = product_map[sku]
        row = [sku, product[1], product[2], copy[2], copy[3], copy[5], copy[4], product[6], copy[6], copy[10], copy[7]]
        for col, value in enumerate(row):
            if col == 7:
                intro.write(4 + offset, col, value, formats["money"])
            elif col in {0, 1, 9}:
                intro.write(4 + offset, col, value, formats["center"])
            else:
                intro.write(4 + offset, col, value, formats["text"])
        intro.set_row(4 + offset, 72)
    intro.add_table(
        3,
        0,
        3 + len(PRODUCT_COPY),
        len(intro_headers) - 1,
        {"name": "SimpleProductCopy", "style": "Table Style Medium 4", "columns": [{"header": h} for h in intro_headers]},
    )
    intro.conditional_format(4, 9, 3 + len(PRODUCT_COPY), 9, {"type": "text", "criteria": "containing", "value": "待", "format": workbook.add_format({"bg_color": PALE_GOLD, "font_color": NAVY})})
    intro.print_area(0, 0, 3 + len(PRODUCT_COPY), len(intro_headers) - 1)

    workbook.close()


def set_cell_shading(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, top: int = 100, start: int = 110, bottom: int = 100, end: int = 110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        set_cell_shading(cell, TEAL)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(9)
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for col_index, value in enumerate(row):
            cells[col_index].text = str(value)
            set_cell_shading(cells[col_index], WHITE if row_index % 2 == 0 else LIGHT)
            set_cell_margins(cells[col_index])
            cells[col_index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cells[col_index].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Cm(width)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.add_run(item)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.3)
    section.left_margin = Cm(1.4)
    section.right_margin = Cm(1.4)
    styles = doc.styles
    styles["Normal"].font.name = "Microsoft JhengHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    styles["Normal"].font.size = Pt(10)
    for style_name, size, color in (("Title", 24, NAVY), ("Heading 1", 16, NAVY), ("Heading 2", 12, TEAL)):
        style = styles[style_name]
        style.font.name = "Microsoft JhengHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("希望之光 Hope Light｜下一步執行與 Moment 內容規劃｜")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)


def _build_detailed_moment_plan() -> None:
    doc = Document()
    configure_document(doc)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("希望之光\n下一步待辦與 Hope Light Moment 內容規劃")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("產品走進生活，故事留下信任｜2026-08-09｜[Codex]")
    run.font.color.rgb = RGBColor.from_string(TEAL)
    run.bold = True

    box = doc.add_table(rows=1, cols=1)
    box.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = box.cell(0, 0)
    set_cell_shading(cell, MINT)
    set_cell_margins(cell, 180, 220, 180, 220)
    cell.text = (
        "這個帳號不是另一個商品目錄，而是記錄『人在什麼時刻需要停下來、如何使用產品、使用後做了什麼』。"
        "內容先建立共鳴與信任，再把有需要的人帶到 LINE 了解蠟燭、大小貴人或適合的服務。"
    )

    doc.add_heading("一、接下來要做的待辦清單", level=1)
    todo_rows = [
        ["P0", "本週", "確認大小貴人價格、版本與差異", "完成可對外使用的版本／價格表；衝突未關閉前不公開報價", "老師＋規劃者", "□"],
        ["P0", "本週", "完成實體產品盤點", "每個SKU有實際數量、盤點日期、預留量與安全庫存", "老師／庫存人員", "□"],
        ["P0", "本週", "補齊蠟燭與大小貴人產品事實", "名稱、規格、使用方式、安全注意、保固與不可宣稱內容都有核准版", "老師", "□"],
        ["P0", "本週", "確認故事與影像同意流程", "服務同意、行銷聯繫、匿名故事公開分開記錄", "老師＋規劃者", "□"],
        ["P0", "本週", "完成 Moment 帳號門面", "名稱、Bio、LINE連結、6個精選動態與3篇置頂內容完成", "規劃者", "□"],
        ["P1", "第2週", "建立一次拍攝素材庫", "完成老師、產品、使用、包裝、空間與手部動作等至少30段直式素材", "規劃者＋拍攝", "□"],
        ["P1", "第2週", "完成首批6篇內容", "4支Reels＋2篇輪播完成審稿、字幕、封面與CTA", "規劃者", "□"],
        ["P1", "第2週", "建立LINE承接SOP", "蠟燭／大小貴人／故事回覆各有關鍵字、問題、負責人與下一步", "老師＋規劃者", "□"],
        ["P1", "第2週", "建立客戶故事採集表", "能記錄情境、使用過程、具體改變、原話、產品、時間與公開同意", "規劃者", "□"],
        ["P2", "第3-4週", "執行4週內容節奏", "每週3篇Feed、4-6天Stories；不因臨時促銷中斷內容主題", "規劃者", "□"],
        ["P2", "每週", "做30分鐘數據回顧", "追蹤非粉絲觸及、觀看留存、收藏、分享、個人檔案造訪、LINE詢問與成交", "規劃者＋老師", "□"],
        ["P2", "第4週", "保留、停做與放大", "選出2個高共鳴題材、1個高轉換題材與1個應停止格式，形成下月計畫", "規劃者＋老師", "□"],
    ]
    add_table(doc, ["優先", "時程", "任務", "完成標準", "負責", "狀態"], todo_rows, [1.1, 1.5, 3.2, 7.0, 2.2, 0.9])

    doc.add_heading("二、Hope Light Moment 帳號角色", level=1)
    add_table(
        doc,
        ["帳號", "主要角色", "內容重點", "主要下一步"],
        [
            ["hopelight.ig", "專業與『看見』", "腦意識觀察、老師方法、課程與深度知識", "加入LINE／預約觀察或課程"],
            ["hopelight.moment", "生活與『陪伴』", "蠟燭、大小貴人、使用情境、老師日常與真實故事", "留言關鍵字／加入LINE了解產品"],
        ],
        [3.0, 3.0, 7.1, 4.0],
    )
    doc.add_heading("2.1 建議帳號門面", level=2)
    add_bullets(
        doc,
        [
            "顯示名稱：Hope Light Moment｜五分鐘生活整理。",
            "Bio第一行：把需要停下來的時刻，變成一個小小儀式。",
            "Bio第二行：蠟燭・大小貴人・真實使用故事。",
            "Bio第三行：不保證結果，只陪你整理當下與下一步。",
            "連結CTA：加入LINE，了解適合自己的產品與使用方式。",
            "精選動態：第一次來／蠟燭／大小貴人／真實故事／使用方法／常見問題。",
            "三篇置頂：帳號在做什麼、十款蠟燭怎麼選、大小貴人版本與使用情境（價格與差異核准後才發布）。",
        ],
    )

    doc.add_heading("三、每週應發多少內容", level=1)
    cadence_rows = [
        ["Feed", "每週3篇", "2支Reels＋1篇輪播", "先連續4週，共12篇；穩定比短期爆量重要"],
        ["Stories", "每週4-6天", "每天3-6格", "幕後、使用片段、投票、問答、客戶原話與導流"],
        ["直播／Q&A", "每月1次", "20-30分鐘", "集中回答選品、使用與安全問題；可拆成短片"],
        ["Trial Reels", "每2週1支", "測試新題材", "先給非粉絲看；24小時後看數據，再決定是否公開給所有人"],
    ]
    add_table(doc, ["形式", "起始頻率", "內容組合", "執行原則"], cadence_rows, [2.5, 2.5, 4.2, 8.0])
    paragraph = doc.add_paragraph()
    paragraph.add_run("這是起始節奏，不是平台保證公式。").bold = True
    paragraph.add_run(" 第4週依專業儀表板的個人化建議與實際留存、分享、LINE詢問調整；若拍攝量不足，寧可維持每週2篇，也不要用重複促銷填滿版面。")

    doc.add_heading("四、內容比例與故事系統", level=1)
    add_table(
        doc,
        ["內容支柱", "比例", "要回答的問題", "適合形式"],
        [
            ["使用後的真實故事", "35%", "這個人原本在哪個情境？使用時做了什麼？之後採取了什麼具體行動？", "Reels／故事輪播"],
            ["生活情境與五分鐘儀式", "30%", "工作、關係、睡前或空間混亂時，可以如何停下來整理？", "Reels／Stories"],
            ["產品認識與選擇", "20%", "蠟燭與大小貴人是什麼、怎麼選、怎麼用、有哪些限制？", "輪播／示範Reels"],
            ["老師與製作幕後", "15%", "老師如何選擇、準備、說明產品？品牌為何重視陪伴與界線？", "口述Reels／幕後Stories"],
        ],
        [3.1, 1.4, 8.0, 4.0],
    )
    add_bullets(
        doc,
        [
            "故事固定記錄七欄：使用前情境、選擇原因、使用方式、當下感受、之後的具體行動、客戶原話、公開同意。",
            "故事不寫成『點了就成功』『用了就改善』；改寫為本人可描述的感受、選擇與行動。",
            "客戶姓名、畫面、訊息截圖與可辨識細節，沒有書面同意就匿名、模糊或不用。",
        ],
    )

    doc.add_heading("五、影片怎麼拍", level=1)
    add_table(
        doc,
        ["段落", "秒數", "畫面", "說話／字幕任務"],
        [
            ["鉤子", "0-2秒", "人臉、手部動作或產品近景", "直接說出一個生活時刻：『今天腦子很滿，我沒有再逼自己想答案。』"],
            ["情境", "3-8秒", "桌面、回家、工作前、睡前", "讓觀眾認出自己，不先講產品名稱"],
            ["使用", "9-25秒", "點燭、選擇、擺放、老師示範", "只示範一個動作與一個原因"],
            ["故事", "26-35秒", "日記、客戶匿名原話、老師旁白", "說具體感受或行動，不保證結果"],
            ["CTA", "36-45秒", "產品＋LINE或留言關鍵字", "每支只要一個下一步：留言、收藏或加入LINE"],
        ],
        [2.0, 1.6, 5.4, 8.0],
    )
    add_bullets(
        doc,
        [
            "以9:16直式、1080×1920拍攝；起始片長20-45秒，一支只講一件事。",
            "固定自然光、乾淨背景與同一組品牌色；至少一半影片出現真人聲音、手或臉，避免只拍靜物廣告。",
            "先錄原音與老師口述，再補字幕；字幕每行短、重點詞上色，封面只放一個問題。",
            "每兩週集中拍2小時：產品特寫、點燭、拿取、充電、包裝、老師說話、書寫、桌面與空間等30段素材。",
            "每支影片保留原檔、內容ID、題材、發布日與CTA，之後才能對照詢問與成交。",
        ],
    )

    doc.add_heading("六、10個文案與影片方向", level=1)
    content_rows = [
        ["1", "五分鐘，不急著找答案", "今天不是需要更多答案，而是先停五分鐘。", "回家放包包→整理桌面→點一顆蠟燭→寫下一件事", "蠟燭生活入口", "留言『五分鐘』"],
        ["2", "十款蠟燭怎麼選", "十款蠟燭不是十個願望，而是十種整理生活問題的入口。", "十款依工作／關係／學習／睡前分組，逐款快速帶過", "十款蠟燭", "收藏選擇圖"],
        ["3", "沒有突然變幸運的故事", "她沒有突然變幸運，只是終於做了拖很久的那個決定。", "匿名情境→使用片段→客戶原話→實際做出的下一步", "蠟燭＋故事", "分享給正在卡住的人"],
        ["4", "大小貴人差異", "大小貴人到底差在哪裡？先不談功效，先看你會在哪裡使用。", "尺寸／頻道／攜帶／空間情境對照；資料核准後拍", "大小貴人", "加入LINE做適配"],
        ["5", "老師自己怎麼用", "老師不是每天都使用同一個頻道，她先問自己這一題。", "老師出鏡說明當下情境→選擇→使用→記錄", "大小貴人", "留言最常使用的時刻"],
        ["6", "新手第一次使用", "包裹打開後，不用急著全部理解，先做這三步。", "開箱→安全確認→第一次使用→如何收納", "蠟燭／大小貴人", "收藏新手步驟"],
        ["7", "客戶留下的一句話", "她用完後沒有說『一切都好了』，只留下這一句。", "匿名訊息局部→老師讀出→解釋為何這句話重要", "產品故事", "留言自己的那一句"],
        ["8", "一個很亂的晚上", "那天晚上，空間很亂、腦子也很滿，我只做了三件事。", "關手機→點燭／放置工具→寫下明天第一步", "蠟燭生活儀式", "收藏晚間流程"],
        ["9", "你今天卡在哪個Moment", "工作、關係、學習、睡前——今天的你最需要整理哪一個？", "四個快速情境＋投票貼紙／留言選項", "產品選擇入口", "留言1／2／3／4"],
        ["10", "七天使用日記", "產品沒有替我做決定，但七天後，我發現它一直提醒我一件事。", "Day1、Day3、Day7固定角度＋三句日記", "蠟燭／大小貴人", "追蹤下一集"],
    ]
    add_table(doc, ["#", "方向", "開場文案", "建議畫面／內容", "產品連結", "CTA"], content_rows, [0.7, 2.4, 5.4, 6.4, 2.3, 2.0])

    doc.add_heading("七、第一個月12篇排程", level=1)
    schedule_rows = [
        ["第1週", "輪播：Moment帳號是什麼", "Reels：五分鐘，不急著找答案", "Reels：老師自己怎麼用"],
        ["第2週", "輪播：十款蠟燭情境地圖", "Reels：沒有突然變幸運的故事", "Reels：大小貴人使用情境差異"],
        ["第3週", "Reels：新手第一次使用", "Reels：一個很亂的晚上", "輪播：產品常見問題與安全界線"],
        ["第4週", "Reels：你今天卡在哪個Moment", "Reels：七天使用日記", "輪播：本月三個真實Moment"],
    ]
    add_table(doc, ["週次", "第1篇", "第2篇", "第3篇"], schedule_rows, [2.0, 5.2, 5.2, 5.2])

    doc.add_heading("八、每週數據只看四層", level=1)
    add_numbered(
        doc,
        [
            "看見：觸及人數、非粉絲觸及、Reels播放。",
            "停留：平均觀看時間、留存曲線、輪播滑到最後一張的訊號。",
            "信任：收藏、分享、留言、私訊與Stories回覆。",
            "行動：個人檔案造訪、連結點擊、LINE新增、產品詢問、報價與成交。",
        ],
    )
    paragraph = doc.add_paragraph()
    paragraph.add_run("判斷原則：").bold = True
    paragraph.add_run("觸及高但沒詢問，補強產品連結與CTA；收藏分享高，延伸成系列；觀看前3秒大量流失，重拍開場；詢問多但不成交，檢查價格、產品差異、信任證據與LINE承接。")

    doc.add_heading("九、Instagram官方原則與本案採用方式", level=1)
    add_bullets(
        doc,
        [
            "Instagram專業儀表板的Best Practices會依帳號提供發布頻率、注意力、Reels長度、互動與觸及的個人化建議；本案先用4週基準，再依帳號實際建議調整。",
            "Trial Reels可先向非粉絲測試新題材，約24小時後查看觀看、按讚、留言與分享；適合測試故事型內容是否能接觸新受眾。",
            "Reels Insights包含重播與留存曲線等指標，因此不能只看播放數，應同時看觀眾在哪裡離開。",
        ],
    )
    sources = doc.add_paragraph()
    sources.add_run("官方來源：").bold = True
    sources.add_run(
        " Meta〈Introducing Best Practices, an Education Hub for Creators on Instagram〉（2024-10-01）；"
        "Meta〈Test Content With Non-Followers Using Trial Reels〉（2024-12-10）；"
        "Meta〈New Ways to Create Content on Instagram〉（2023-11-15）。"
    )

    doc.save(MOMENT_DOCX)


def build_simple_workbook() -> None:
    """Build the one-person-studio version with only fields used every week."""
    workbook = xlsxwriter.Workbook(SIMPLE_XLSX)
    workbook.set_properties(
        {
            "title": "希望之光一人工作室簡易產品表",
            "subject": "報價、盤點與產品說法",
            "author": "[Codex]",
            "company": "希望之光 Hope Light",
            "created": date(2026, 8, 9),
        }
    )
    formats = excel_formats(workbook)

    quote_headers = ["產品／服務", "單價", "價格狀態", "數量", "小計", "備註"]
    quote_rows = [
        ["頻率蠟燭1-10號（任選）", 139, "來源售價，待老師核准", "", "", "報價時填寫蠟燭號碼"],
        ["21顆蠟燭＋10分鐘諮詢首購組", 2980, "來源售價，待老師核准", "", "", "內容、運費與成本待確認"],
        ["大貴人－9個頻道", 16800, "價格衝突，暫勿對外", "", "", "來源另有24,900"],
        ["小貴人－9個頻道", 24900, "價格衝突，暫勿對外", "", "", "來源另有16,800"],
        ["大貴人－7個頻道", 16800, "價格衝突，暫勿對外", "", "", "版本差異待確認"],
        ["腦意識觀察／檢測90分鐘", 2980, "來源售價，待老師核准", "", "", "正式名稱待老師核准"],
        ["深度諮詢60分鐘", 1980, "已確認，可報價", "", "", "[User] 2026-08-09確認"],
    ]
    quote = workbook.add_worksheet("簡易報價")
    setup_sheet(
        quote,
        formats,
        "希望之光｜一人工作室簡易報價",
        "平常只用這7項。黃底填數量；價格有衝突的大小貴人先不要對外報價。金額單位：新台幣。",
        quote_headers,
        [36, 13, 24, 10, 14, 38],
    )
    for offset, row in enumerate(quote_rows):
        sheet_row = 4 + offset
        excel_row = sheet_row + 1
        for col, value in enumerate(row):
            if col == 1:
                quote.write(sheet_row, col, value, formats["money"])
            elif col == 3:
                quote.write_blank(sheet_row, col, None, formats["input"])
            elif col == 4:
                quote.write_formula(
                    sheet_row,
                    col,
                    f'=IF(OR(B{excel_row}="",D{excel_row}=""),"",B{excel_row}*D{excel_row})',
                    formats["money"],
                    "",
                )
            elif col == 2:
                quote.write(sheet_row, col, value, formats["center"])
            else:
                quote.write(sheet_row, col, value, formats["text"])
        quote.set_row(sheet_row, 40)
    quote.add_table(
        3,
        0,
        3 + len(quote_rows),
        len(quote_headers) - 1,
        {"name": "SoloQuote", "style": "Table Style Medium 4", "columns": [{"header": h} for h in quote_headers]},
    )
    quote.data_validation(4, 3, 3 + len(quote_rows), 3, {"validate": "integer", "criteria": ">=", "value": 0})
    quote.conditional_format(4, 2, 3 + len(quote_rows), 2, {"type": "text", "criteria": "containing", "value": "已確認", "format": workbook.add_format({"bg_color": PALE_GREEN, "font_color": GREEN})})
    quote.conditional_format(4, 2, 3 + len(quote_rows), 2, {"type": "text", "criteria": "containing", "value": "衝突", "format": workbook.add_format({"bg_color": PALE_RED, "font_color": RED})})

    inventory_headers = ["產品", "盤點數量", "本週售出", "剩餘數量", "盤點日期", "備註"]
    inventory_names = [
        "1號 防護除穢",
        "2號 財富豐盛",
        "3號 吸引顧客",
        "4號 吸引桃花",
        "5號 貴人常臨",
        "6號 學業進步",
        "7號 小人退散",
        "8號 好運爆棚",
        "9號 感情升溫",
        "10號 舒壓好眠",
        "大貴人－9個頻道",
        "小貴人－9個頻道",
        "大貴人－7個頻道",
    ]
    inventory = workbook.add_worksheet("簡易庫存")
    setup_sheet(
        inventory,
        formats,
        "希望之光｜一人工作室簡易庫存",
        "每週固定一天盤點即可。只填黃底三欄，不使用舊表數字；剩餘數量會自動計算。",
        inventory_headers,
        [36, 13, 13, 13, 15, 42],
    )
    for offset, name in enumerate(inventory_names):
        sheet_row = 4 + offset
        excel_row = sheet_row + 1
        inventory.write(sheet_row, 0, name, formats["text"])
        inventory.write_blank(sheet_row, 1, None, formats["input"])
        inventory.write_blank(sheet_row, 2, None, formats["input"])
        inventory.write_formula(
            sheet_row,
            3,
            f'=IF(B{excel_row}="","",B{excel_row}-IF(C{excel_row}="",0,C{excel_row}))',
            formats["center"],
            "",
        )
        inventory.write_blank(sheet_row, 4, None, formats["date_input"])
        inventory.write_blank(sheet_row, 5, None, formats["text"])
        inventory.set_row(sheet_row, 38)
    inventory.add_table(
        3,
        0,
        3 + len(inventory_names),
        len(inventory_headers) - 1,
        {"name": "SoloInventory", "style": "Table Style Medium 4", "columns": [{"header": h} for h in inventory_headers]},
    )
    inventory.data_validation(4, 1, 3 + len(inventory_names), 2, {"validate": "integer", "criteria": ">=", "value": 0})
    inventory.data_validation(4, 4, 3 + len(inventory_names), 4, {"validate": "date", "criteria": "between", "minimum": date(2020, 1, 1), "maximum": date(2100, 12, 31)})

    intro_headers = ["產品／服務", "一句話怎麼說", "適合什麼時刻", "客戶下一步", "目前狀態"]
    intro_rows = [
        ["頻率蠟燭1-10號", "把需要停下來的時刻，變成五分鐘整理自己的小儀式。", "工作、關係、學習、回家後或睡前", "選一個現在最需要整理的情境", "名稱與內容待老師核准"],
        ["大貴人－9個頻道", "在需要安靜整理狀態與行動節奏時，提供固定的陪伴提醒。", "會談、工作轉換或個人整理前", "加入LINE了解頻道與使用方式", "價格、頻道與規格待確認"],
        ["小貴人－9個頻道", "把短暫停頓變成日常可以重複的節奏提醒。", "會議、專注、休息或外出時", "加入LINE了解適合情境", "價格、頻道與規格待確認"],
        ["大貴人－7個頻道", "較精簡的日常整理版本，先依使用情境確認是否適合。", "待老師補充", "加入LINE詢問版本差異", "版本差異與價格待確認"],
        ["腦意識觀察／檢測90分鐘", "用結構化觀察與對談，整理自己的思維模式與下一步。", "反覆卡在工作、關係或選擇時", "加入LINE預約", "售價與正式名稱待核准"],
        ["深度諮詢60分鐘", "把零散的感受與問題，整理成一個可以前進的方向。", "已看見問題、需要進一步整理時", "加入LINE選擇諮詢主題", "售價1,980已確認"],
    ]
    intro = workbook.add_worksheet("簡易介紹")
    setup_sheet(
        intro,
        formats,
        "希望之光｜一人工作室簡易產品介紹",
        "老師不需要背長文，只要先用『一句話怎麼說』回答。避免治療、改運與保證效果。",
        intro_headers,
        [30, 58, 40, 34, 30],
    )
    for offset, row in enumerate(intro_rows):
        for col, value in enumerate(row):
            intro.write(4 + offset, col, value, formats["text"] if col != 4 else formats["center"])
        intro.set_row(4 + offset, 58)
    intro.add_table(
        3,
        0,
        3 + len(intro_rows),
        len(intro_headers) - 1,
        {"name": "SoloIntro", "style": "Table Style Medium 4", "columns": [{"header": h} for h in intro_headers]},
    )
    workbook.close()


def build_moment_plan() -> None:
    """Build a low-burden content plan for a one-person studio."""
    doc = Document()
    configure_document(doc)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("希望之光一人工作室\nHope Light Moment 簡單執行表")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("每週只做：拍一次・發兩篇・記一次詢問與訂單｜2026-08-09｜[Codex]").bold = True

    box = doc.add_table(rows=1, cols=1)
    box.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = box.cell(0, 0)
    set_cell_shading(cell, MINT)
    set_cell_margins(cell, 180, 220, 180, 220)
    cell.text = "老師是一人工作室，不需要像公司一樣每天發文、維護複雜CRM或追很多數字。先讓客戶看懂蠟燭與大小貴人，再用真實故事累積信任。"

    doc.add_heading("一、先完成這6件事", level=1)
    todo_rows = [
        ["1", "確認主力品項", "先只推蠟燭、大小貴人與60分鐘諮詢，不同時宣傳全部產品", "本週"],
        ["2", "確認價格", "先關閉大小貴人價格衝突；諮詢價固定1,980", "本週"],
        ["3", "做一次盤點", "把13項實體品項的現場數量填進簡易庫存表", "本週"],
        ["4", "整理帳號門面", "完成Bio、LINE連結與3篇置頂內容", "第1週"],
        ["5", "集中拍1小時", "一次錄6-10段產品、老師與使用畫面，之後慢慢剪", "每2週"],
        ["6", "固定每週記錄", "只記LINE詢問數、訂單數、本週最有效的一篇", "每週10分鐘"],
    ]
    add_table(doc, ["#", "要做什麼", "做到什麼就算完成", "時間"], todo_rows, [0.8, 3.6, 10.8, 2.4])

    doc.add_heading("二、兩個IG帳號不要同時重度經營", level=1)
    add_table(
        doc,
        ["帳號", "現在怎麼用", "頻率"],
        [
            ["hopelight.moment", "主帳號：蠟燭、大小貴人、日常使用與真實故事", "每週2篇＋Stories每週3天"],
            ["hopelight.ig", "專業資料庫：腦意識觀察、老師方法與課程；可沿用Moment素材", "每月1篇即可，忙時可暫停"],
        ],
        [4.0, 9.5, 4.2],
    )
    add_bullets(
        doc,
        [
            "Moment顯示名稱：Hope Light Moment｜五分鐘生活整理。",
            "Bio：把需要停下來的時刻，變成一個小小儀式｜蠟燭・大小貴人・真實故事｜加入LINE了解適合自己的選擇。",
            "置頂3篇：這個帳號在做什麼／十款蠟燭怎麼選／大小貴人怎麼選（價格與版本確認後再發）。",
            "精選動態只留4個：蠟燭／大小貴人／故事／怎麼買。",
        ],
    )

    doc.add_heading("三、老師每週只需要60-90分鐘", level=1)
    add_table(
        doc,
        ["誰", "每週要做的事", "大約時間"],
        [
            ["老師", "提供1個故事、回答1個產品問題、集中出鏡或拍手部畫面", "60-90分鐘"],
            ["品牌規劃者", "整理題目、寫短稿、剪輯、封面、排程與簡單記錄", "依合作安排"],
        ],
        [3.2, 10.5, 4.0],
    )
    doc.add_heading("3.1 最小每週節奏", level=2)
    add_table(
        doc,
        ["時間", "動作", "內容"],
        [
            ["週一15分鐘", "選題", "從10個方向挑1個產品題、1個故事題"],
            ["任一天45-60分鐘", "集中拍攝", "一次拍6-10個短畫面，不當天剪"],
            ["週三", "發1支Reels", "15-30秒產品／使用情境"],
            ["週六", "發1篇照片或輪播", "故事、客戶一句話或選品說明"],
            ["每週3天", "Stories", "每次1-3格：幕後、投票、問題或出貨"],
            ["週日10分鐘", "簡單記錄", "LINE詢問、訂單、本週最好的一篇"],
        ],
        [3.2, 3.4, 11.0],
    )

    doc.add_heading("四、內容只分三類", level=1)
    add_table(
        doc,
        ["內容", "比例", "怎麼拍"],
        [
            ["使用情境", "40%", "工作前、回家後、睡前、關係卡住時，示範一個簡單動作"],
            ["真實故事", "40%", "使用前的情境、使用時做了什麼、之後採取什麼行動"],
            ["產品介紹", "20%", "怎麼選、怎麼用、差異、價格與注意事項"],
        ],
        [4.0, 2.0, 11.6],
    )
    paragraph = doc.add_paragraph()
    paragraph.add_run("影片只用三個畫面：").bold = True
    paragraph.add_run("①現在遇到什麼時刻 → ②怎麼使用產品 → ③一句真實感受或下一步。片長15-30秒，最後只放一個CTA。")

    doc.add_heading("五、10個可以直接拍的方向", level=1)
    content_rows = [
        ["1", "今天腦子很滿，先停五分鐘", "桌面很亂→點燭→寫下一件事", "留言『五分鐘』"],
        ["2", "十款蠟燭怎麼選", "十款排開→分成工作／關係／學習／睡前", "收藏這張選擇表"],
        ["3", "她沒有突然變幸運", "匿名故事→產品片段→她實際做出的決定", "分享給卡住的人"],
        ["4", "大小貴人怎麼選", "大小與使用情境對照；資料確認後才發布", "加入LINE詢問"],
        ["5", "老師自己什麼時候使用", "老師一句話→拿取產品→使用片段", "留言你的使用時刻"],
        ["6", "第一次使用只做三步", "開箱→安全確認→開始使用", "收藏新手步驟"],
        ["7", "客戶留下的一句話", "匿名原話→老師讀出→一個具體行動", "留言你的那一句"],
        ["8", "一個很亂的晚上", "關手機→點燭→寫明天第一步", "收藏晚間流程"],
        ["9", "今天卡在哪個Moment", "工作／關係／學習／睡前四選一", "留言1、2、3、4"],
        ["10", "七天簡單使用日記", "Day1／Day3／Day7各一個固定畫面", "追蹤下一集"],
    ]
    add_table(doc, ["#", "題目", "三個畫面", "CTA"], content_rows, [0.8, 5.0, 8.6, 3.2])

    doc.add_heading("六、第一個月只發8篇", level=1)
    add_table(
        doc,
        ["週次", "Reels", "照片／輪播"],
        [
            ["第1週", "今天腦子很滿，先停五分鐘", "Moment帳號在做什麼"],
            ["第2週", "老師自己什麼時候使用", "十款蠟燭怎麼選"],
            ["第3週", "第一次使用只做三步", "一個匿名客戶故事"],
            ["第4週", "今天卡在哪個Moment", "本月最常被問的產品問題"],
        ],
        [2.2, 7.5, 7.5],
    )

    doc.add_heading("七、每週只看三個數字", level=1)
    add_numbered(
        doc,
        [
            "本週有幾個LINE產品詢問？",
            "本週有幾張訂單？",
            "哪一篇最容易被收藏、分享或收到回覆？下週再做一篇相近題目。",
        ],
    )
    paragraph = doc.add_paragraph()
    paragraph.add_run("不要做的事：").bold = True
    paragraph.add_run("不每天發文、不同時推所有產品、不維護複雜CRM、不追十幾個KPI、不把故事寫成保證功效。")
    doc.save(MOMENT_DOCX)


def configure_solo_roadmap_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)
    styles = doc.styles
    styles["Normal"].font.name = "Microsoft JhengHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    styles["Normal"].font.size = Pt(10.5)
    for style_name, size, color in (("Title", 24, NAVY), ("Heading 1", 16, NAVY), ("Heading 2", 12, TEAL)):
        style = styles[style_name]
        style.font.name = "Microsoft JhengHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("希望之光｜一人工作室簡易品牌路徑｜")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)


def build_solo_roadmap(path: Path = ROADMAP_DOCX) -> None:
    """Replace the company-style roadmap with a one-person-studio version."""
    doc = Document()
    configure_solo_roadmap_document(doc)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("希望之光\n一人工作室簡易品牌路徑")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("先讓工作變簡單，再讓客戶慢慢增加｜2026-08-09｜[Codex]").bold = True

    box = doc.add_table(rows=1, cols=1)
    cell = box.cell(0, 0)
    set_cell_shading(cell, MINT)
    set_cell_margins(cell, 180, 220, 180, 220)
    cell.text = "希望之光現階段不是要建立一間大公司，而是讓老師一個人也能穩定完成：客戶看得懂、產品找得到、有人詢問時回得出來、成交後記得回訪。"

    doc.add_heading("一、只守住四個原則", level=1)
    add_numbered(
        doc,
        [
            "一次只推兩類主力：頻率蠟燭與大小貴人；需要深度整理時再承接60分鐘諮詢。",
            "一週只發兩篇內容，不追求每天更新。",
            "所有價格、庫存與一句話介紹，只維護一份簡易Excel。",
            "每週只看LINE詢問、訂單、本週最好的一篇，不建立複雜報表。",
        ],
    )

    doc.add_heading("二、產品怎麼排優先順序", level=1)
    add_table(
        doc,
        ["優先", "產品／服務", "角色", "現在要做的事"],
        [
            ["1", "頻率蠟燭", "最容易開始的生活產品", "把十款用途、價格、庫存與安全說明整理清楚"],
            ["2", "大小貴人", "需要說明與信任的核心產品", "先確認版本、價格、頻道與怎麼使用，再拍內容"],
            ["3", "60分鐘諮詢", "產品之外的深度承接", "固定售價1,980，說清楚三種工具差異"],
            ["暫緩", "其他設備與課程", "未來選項", "資料未完成前不主動推廣，避免老師同時承接太多"],
        ],
        [1.6, 3.8, 5.0, 7.0],
    )

    doc.add_heading("三、30天只走四步", level=1)
    add_table(
        doc,
        ["週次", "只做一件大事", "完成標準"],
        [
            ["第1週", "把主力產品說清楚", "價格、名稱、庫存、使用方式與注意事項可回答"],
            ["第2週", "把Moment帳號整理好", "Bio、LINE連結、4個精選與3篇置頂完成"],
            ["第3週", "開始每週兩篇", "1支短影片＋1篇照片／輪播，Stories每週3天"],
            ["第4週", "看詢問與訂單", "留下有效題材，刪掉沒人理解的說法，再排下月8篇"],
        ],
        [2.3, 6.0, 9.2],
    )

    doc.add_heading("四、最簡單的客戶流程", level=1)
    flow = doc.add_paragraph()
    flow.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = flow.add_run("看到貼文／Stories  →  加入LINE  →  回答三個問題  →  推薦一項產品  →  成交  →  7天後簡單回訪")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(TEAL)
    add_bullets(
        doc,
        [
            "LINE只問三題：現在想整理什麼情境？比較想了解蠟燭還是大小貴人？希望自己使用還是送人？",
            "一次只推薦一項產品，說明原因、價格與使用方式，不一次丟出全部商品。",
            "成交後7天只問一句：『這幾天你在哪個時刻使用？有沒有一個感受或行動想記下來？』",
            "客戶故事要公開前，再取得文字、照片或匿名使用同意。",
        ],
    )

    doc.add_heading("五、老師與品牌規劃者怎麼分工", level=1)
    add_table(
        doc,
        ["角色", "只負責什麼", "不用負責什麼"],
        [
            ["老師", "確認產品事實、每週提供1個故事、集中出鏡60-90分鐘、回覆重要LINE", "腳本排版、剪片、複雜數據與每天發文"],
            ["品牌規劃者", "把老師的話整理成短稿、拍攝清單、剪輯、排程與每週簡單紀錄", "替老師保證銷量或同時放大所有產品"],
        ],
        [3.0, 7.3, 7.3],
    )

    doc.add_heading("六、一週最小工作量", level=1)
    add_table(
        doc,
        ["動作", "頻率", "大約時間"],
        [
            ["確認本週兩個題目", "每週一次", "15分鐘"],
            ["集中拍攝6-10段畫面", "每兩週一次", "45-60分鐘"],
            ["發布內容", "每週兩篇", "由規劃者整理"],
            ["Stories", "每週三天，每次1-3格", "有素材才發"],
            ["記錄詢問、訂單與最佳貼文", "每週一次", "10分鐘"],
        ],
        [6.8, 6.0, 4.8],
    )

    doc.add_heading("七、現在先不要做", level=1)
    add_bullets(
        doc,
        [
            "不要同時重度經營兩個IG帳號；Moment先做主帳號，專業帳號每月一篇或暫停。",
            "不要一次推全部23個品項；其他設備與課程等資料完整後再排。",
            "不要建立老師每天都要填的CRM與KPI表；先用LINE、簡易產品表與每週三個數字。",
            "不要要求每天拍片；每兩週集中拍一次，素材分批使用。",
            "不要使用治療、診斷、改運或保證結果的文案。",
        ],
    )
    doc.add_heading("八、下一步", level=1)
    add_numbered(
        doc,
        [
            "老師確認大小貴人版本與價格。",
            "把現場庫存填入簡易庫存表。",
            "完成Moment帳號Bio、LINE連結與第一個月8篇內容。",
        ],
    )
    doc.save(path)


def main() -> None:
    build_simple_workbook()
    build_moment_plan()
    build_solo_roadmap()
    print(SIMPLE_XLSX)
    print(MOMENT_DOCX)
    print(ROADMAP_DOCX)


if __name__ == "__main__":
    main()
