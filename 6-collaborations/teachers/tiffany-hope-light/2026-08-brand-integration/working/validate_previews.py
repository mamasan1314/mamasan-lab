from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

HERE = Path(__file__).resolve().parent
sys.path.insert(0, str(HERE / "vendor"))

import pymupdf as fitz
from PIL import Image, ImageDraw
from docx import Document


def render_pdf(pdf_path: Path, output_dir: Path, max_pages: int | None = None) -> dict[str, object]:
    doc = fitz.open(pdf_path)
    pages = min(len(doc), max_pages) if max_pages else len(doc)
    thumbs: list[Image.Image] = []
    text_lengths: list[int] = []
    for index in range(pages):
        page = doc[index]
        text_lengths.append(len(page.get_text("text").strip()))
        pix = page.get_pixmap(matrix=fitz.Matrix(0.8, 0.8), alpha=False)
        image = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
        image.thumbnail((300, 400))
        tile = Image.new("RGB", (320, 440), "white")
        tile.paste(image, ((320 - image.width) // 2, 25))
        ImageDraw.Draw(tile).text((10, 7), f"{pdf_path.stem} p.{index + 1}", fill="black")
        thumbs.append(tile)
    if thumbs:
        columns = 3
        rows = (len(thumbs) + columns - 1) // columns
        contact = Image.new("RGB", (columns * 320, rows * 440), "#d7d7d7")
        for index, thumb in enumerate(thumbs):
            contact.paste(thumb, ((index % columns) * 320, (index // columns) * 440))
        contact_path = output_dir / f"{pdf_path.stem}_contact.jpg"
        contact.save(contact_path, quality=90)
    else:
        contact_path = None
    result = {
        "file": str(pdf_path),
        "pages": len(doc),
        "rendered_pages": pages,
        "text_characters": sum(text_lengths),
        "blank_text_pages": sum(length == 0 for length in text_lengths),
        "contact_sheet": str(contact_path) if contact_path else "",
    }
    doc.close()
    return result


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("output_dir", type=Path)
    parser.add_argument("docx_path", type=Path)
    parser.add_argument("pdf_path", type=Path)
    args = parser.parse_args()
    args.output_dir.mkdir(parents=True, exist_ok=True)

    word = Document(args.docx_path)
    report: dict[str, object] = {
        "docx": {
            "paragraphs": len(word.paragraphs),
            "tables": len(word.tables),
            "nonempty_paragraphs": sum(bool(p.text.strip()) for p in word.paragraphs),
            "characters": sum(len(p.text) for p in word.paragraphs) + sum(
                len(cell.text) for table in word.tables for row in table.rows for cell in row.cells
            ),
        },
        "roadmap_pdf": render_pdf(args.pdf_path, args.output_dir),
        "excel_previews": [],
    }
    for pdf_path in sorted(args.output_dir.glob("*.pdf")):
        report["excel_previews"].append(render_pdf(pdf_path, args.output_dir, max_pages=1))
    report_path = args.output_dir / "preview_validation.json"
    report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
