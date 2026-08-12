from __future__ import annotations

import json
import sys
from datetime import date
from pathlib import Path

HERE = Path(__file__).resolve().parent
PROJECT = HERE.parent
OUTPUT = PROJECT / "outputs"
VENDOR = HERE / "vendor"
sys.path.insert(0, str(VENDOR))
sys.path.insert(0, str(HERE))

import xlsxwriter
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from openpyxl import load_workbook

from deliverable_data import (
    CONTENT_IDEAS,
    CRM_SEED,
    DATA_ISSUES,
    DECISIONS,
    LINE_SOPS,
    PRODUCTS,
    PRODUCT_COPY,
    ROADMAP,
    SOURCE_AUDIT,
)


NAVY = "16253D"
NAVY_2 = "243B5A"
GOLD = "D3A647"
LIGHT_GOLD = "F4E9CD"
PALE_BLUE = "EAF1F8"
PALE_GREEN = "E8F3EB"
PALE_RED = "FCE8E6"
PALE_YELLOW = "FFF4CC"
TEXT = "25313C"
MID_GREY = "6B7785"
LIGHT_GREY = "EEF1F4"
WHITE = "FFFFFF"
GREEN = "2E7D5B"
RED = "B33A3A"


def pad_rows(seed: list[list[object]], row_count: int, width: int) -> list[list[object]]:
    rows = [list(row) + [""] * (width - len(row)) for row in seed]
    rows.extend([[""] * width for _ in range(max(0, row_count - len(rows)))])
    return rows[:row_count]


class WorkbookBuilder:
    def __init__(self, path: Path) -> None:
        self.path = path
        self.wb = xlsxwriter.Workbook(path)
        self.wb.set_properties(
            {
                "title": "希望之光品牌營運管理系統",
                "subject": "品牌、產品、庫存、CRM、內容與90天執行",
                "author": "[Codex] for Hope Light",
                "company": "希望之光 Hope Light",
                "comments": "依2026-08-09提供來源整理；待確認項目不可視為正式核准資料。",
            }
        )
        self.title = self.wb.add_format(
            {
                "font_name": "Microsoft JhengHei",
                "font_size": 20,
                "bold": True,
                "font_color": WHITE,
                "bg_color": NAVY,
                "align": "left",
                "valign": "vcenter",
            }
        )
        self.subtitle = self.wb.add_format(
            {
                "font_name": "Microsoft JhengHei",
                "font_size": 10,
                "font_color": MID_GREY,
                "text_wrap": True,
                "valign": "top",
            }
        )
        self.section = self.wb.add_format(
            {
                "font_name": "Microsoft JhengHei",
                "font_size": 12,
                "bold": True,
                "font_color": WHITE,
                "bg_color": NAVY_2,
                "align": "left",
                "valign": "vcenter",
                "border": 0,
            }
        )
        self.body = self.wb.add_format(
            {
                "font_name": "Microsoft JhengHei",
                "font_size": 10,
                "font_color": TEXT,
                "text_wrap": True,
                "valign": "top",
            }
        )
        self.input_fmt = self.wb.add_format(
            {
                "font_name": "Microsoft JhengHei",
                "font_size": 10,
                "font_color": TEXT,
                "bg_color": PALE_YELLOW,
                "border": 1,
                "border_color": "D9D9D9",
            }
        )
        self.formula_fmt = self.wb.add_format(
            {
                "font_name": "Microsoft JhengHei",
                "font_size": 10,
                "font_color": TEXT,
                "bg_color": PALE_BLUE,
                "border": 1,
                "border_color": "D9D9D9",
            }
        )
        self.money_fmt = self.wb.add_format(
            {
                "font_name": "Microsoft JhengHei",
                "font_size": 10,
                "num_format": '#,##0;[Red]-#,##0',
            }
        )
        self.percent_fmt = self.wb.add_format(
            {"font_name": "Microsoft JhengHei", "font_size": 10, "num_format": "0.0%"}
        )
        self.date_fmt = self.wb.add_format(
            {"font_name": "Microsoft JhengHei", "font_size": 10, "num_format": "yyyy-mm-dd"}
        )
        self.kpi_label = self.wb.add_format(
            {
                "font_name": "Microsoft JhengHei",
                "font_size": 10,
                "bold": True,
                "font_color": WHITE,
                "bg_color": NAVY_2,
                "align": "center",
                "valign": "vcenter",
                "border": 1,
                "border_color": WHITE,
            }
        )
        self.kpi_value = self.wb.add_format(
            {
                "font_name": "Microsoft JhengHei",
                "font_size": 18,
                "bold": True,
                "font_color": NAVY,
                "bg_color": LIGHT_GOLD,
                "align": "center",
                "valign": "vcenter",
                "border": 1,
                "border_color": GOLD,
            }
        )
        self.table_style = "Table Style Medium 2"

    def add_sheet(self, name: str, purpose: str, tab_color: str = NAVY) -> xlsxwriter.worksheet.Worksheet:
        ws = self.wb.add_worksheet(name)
        ws.set_tab_color(tab_color)
        ws.hide_gridlines(2)
        ws.set_row(0, 34)
        ws.merge_range(0, 0, 0, 12, name.replace("_", "｜", 1), self.title)
        ws.merge_range(1, 0, 1, 12, purpose, self.subtitle)
        ws.set_row(1, 31)
        ws.freeze_panes(4, 0)
        ws.set_default_row(20)
        ws.set_landscape()
        ws.fit_to_pages(1, 0)
        ws.set_margins(0.25, 0.25, 0.4, 0.4)
        ws.set_header("&L希望之光 Hope Light&C品牌營運管理系統&R&D")
        ws.set_footer("&L[Codex] 初版｜待確認資料不可直接對外&R第 &P / &N 頁")
        return ws

    def add_table(
        self,
        ws: xlsxwriter.worksheet.Worksheet,
        start_row: int,
        start_col: int,
        name: str,
        headers: list[str],
        data: list[list[object]],
        formula_columns: dict[str, str] | None = None,
        style: str | None = None,
    ) -> None:
        columns: list[dict[str, object]] = []
        formula_columns = formula_columns or {}
        for header in headers:
            column: dict[str, object] = {"header": header}
            if header in formula_columns:
                column["formula"] = formula_columns[header]
            columns.append(column)
        ws.add_table(
            start_row,
            start_col,
            start_row + len(data),
            start_col + len(headers) - 1,
            {
                "name": name,
                "style": style or self.table_style,
                "columns": columns,
                "data": data,
                "autofilter": True,
            },
        )

    def build_lists(self) -> None:
        ws = self.wb.add_worksheet("19_選單")
        lists = {
            "SKU_LIST": [row[0] for row in PRODUCTS],
            "INV_TYPE_LIST": ["期初", "進貨", "銷售", "退貨入庫", "退貨出庫", "盤盈", "盤虧", "贈品", "樣品"],
            "YES_NO_UNKNOWN": ["是", "否", "未知"],
            "GENERAL_STATUS": ["未開始", "進行中", "待老師", "待資料", "待核准", "已完成", "已停止"],
            "CUSTOMER_TYPE": ["潛在客", "一般客戶", "學生", "經銷／合作夥伴", "轉介紹人", "其他"],
            "CRM_STATUS": ["新名單", "培養中", "已預約", "已成交", "持續服務", "沉睡", "不聯絡", "待補資料"],
            "CHANNEL_LIST": ["IG主帳", "IG Moment", "FB", "官方LINE", "網站", "轉介紹", "現場", "舊客回流", "其他"],
            "FIVE_DIM_LIST": ["事", "財", "情", "家", "體", "多重", "待確認"],
            "PAYMENT_STATUS": ["未請款", "待付款", "部分付款", "已付款", "退款中", "已退款", "取消"],
            "FULFILL_STATUS": ["待確認", "待出貨", "已出貨", "已完成", "退換處理", "取消"],
            "RISK_LEVEL": ["低", "中", "高"],
            "CONTENT_STATUS": ["想法", "待訪談", "待腳本", "待拍攝", "待核准", "已排程", "已發布", "停止"],
            "CONTENT_TYPE": ["真實故事", "FAQ", "知識", "幕後", "見證", "行動"],
            "FUNNEL_STAGE": ["新名單", "已回覆", "有效對話", "已加LINE", "已預約", "已成交", "回購", "暫緩", "流失"],
            "CONSENT_STATUS": ["未知", "不同意", "僅內部", "匿名公開", "具名公開", "已撤回"],
            "PRIVACY_LEVEL": ["一般", "內部", "敏感", "高度敏感"],
            "OWNER_LIST": ["老師", "品牌規劃者", "客服", "出貨", "共同", "待指派"],
            "PLATFORM_LIST": ["Instagram", "Facebook", "LINE", "YouTube", "網站", "其他"],
        }
        for col, (name, values) in enumerate(lists.items()):
            ws.write(0, col, name)
            for row, value in enumerate(values, start=1):
                ws.write(row, col, value)
            col_letter = xlsxwriter.utility.xl_col_to_name(col)
            self.wb.define_name(name, f"='19_選單'!${col_letter}$2:${col_letter}${len(values)+1}")
        ws.hide()

    def build_guide(self) -> None:
        ws = self.add_sheet(
            "00_使用說明",
            "先完成產品與期初盤點，再開始輸入交易。黃色＝人工輸入；藍色＝公式／系統欄；紅色或『待確認』＝不可直接對外使用。",
            GOLD,
        )
        ws.set_column("A:A", 3)
        ws.set_column("B:B", 21)
        ws.set_column("C:C", 70)
        ws.set_column("D:D", 24)
        ws.merge_range("B4:D4", "這份工作簿的角色", self.section)
        ws.write("B5", "品牌作業系統", self.body)
        ws.write(
            "C5",
            "把品牌定位、產品、庫存、訂單、客戶、服務、回訪、內容、故事、LINE與分潤放在同一個可追溯系統中。它不是會計軟體，也不取代正式法務、醫療或個資合規審查。",
            self.body,
        )
        ws.write("D5", "版本：2026-08-09 初版｜[Codex]", self.body)
        ws.merge_range("B7:D7", "第一次使用：依序完成", self.section)
        steps = [
            ("1. 關閉決策", "到 17_品牌決策，把母品牌、價格、產品名稱、宣稱界線等由老師核准。"),
            ("2. 確認產品", "在 03_產品主檔確認SKU、正式名稱、售價、成本與庫存管理；待確認資料不要報價。"),
            ("3. 現場盤點", "所有實體SKU在 05_庫存異動各輸入一筆『期初』，日期使用實際盤點日。"),
            ("4. 建立客戶ID", "新客先進 08_客戶CRM，再用同一客戶ID串接訂單、服務、回訪與故事。"),
            ("5. 每筆都落表", "銷售進 07_銷售訂單；實體數量另進 05_庫存異動。不要再往右新增日期欄。"),
            ("6. 每週驗證", "內容發佈後填 12_內容成效；LINE對話進 14_LINE漏斗；每週只放大可追溯的有效項。"),
            ("7. 故事先同意", "任何個案／見證先進 13_希望故事庫，保留原始證據、匿名處理與公開同意。"),
        ]
        for idx, (label, text) in enumerate(steps, start=8):
            ws.write(idx - 1, 1, label, self.input_fmt)
            ws.merge_range(idx - 1, 2, idx - 1, 3, text, self.body)
        ws.merge_range("B17:D17", "日常操作節奏", self.section)
        cadence = [
            ("每天", "訂單、庫存異動、服務紀錄、LINE下一步在當日完成。"),
            ("每週", "30分鐘看一次：有效對話、LINE新增、預約、訂單、回訪逾期、低庫存。"),
            ("每月", "關閉上月數據、核對實物庫存、檢視毛利與分潤、淘汰無效內容。"),
            ("每季", "重看品牌決策、產品角色、客群與證據；不要只加內容而不修路徑。"),
        ]
        for idx, (label, text) in enumerate(cadence, start=18):
            ws.write(idx - 1, 1, label, self.kpi_label)
            ws.merge_range(idx - 1, 2, idx - 1, 3, text, self.body)
        ws.merge_range("B24:D24", "目前最重要的資料警示", self.section)
        warnings = [
            "大小貴人價格在不同分頁互換，未核准前不可正式報價。",
            "舊庫存表的 #REF!、500件進貨與9/1日期未視為真實資料；需重新盤點。",
            "MEGAWING 2008-2012國外客戶未匯入本CRM，避免舊公司資料污染與擴大個資用途。",
            "諧和機手冊的健康效果用語已列入高風險；本版產品介紹只保留生活情境表述。",
        ]
        for idx, warning in enumerate(warnings, start=25):
            ws.merge_range(idx - 1, 1, idx - 1, 3, f"• {warning}", self.body)
        ws.set_row(1, 38)
        ws.autofit()
        ws.set_column("C:C", 70)

    def build_dashboard(self) -> None:
        ws = self.add_sheet(
            "01_品牌儀表板",
            "開啟Excel後公式會依各工作表自動更新。初版沒有假造業績；數字為0代表尚未輸入，而不是表現不佳。",
            GOLD,
        )
        for col in range(12):
            ws.set_column(col, col, 14)
        blocks = [
            ("B4:C4", "B5:C6", "90天任務完成率", '=IFERROR(COUNTIF(tblRoadmap[狀態],"已完成")/COUNTA(tblRoadmap[任務]),0)', self.percent_fmt),
            ("E4:F4", "E5:F6", "待確認資料", '=COUNTIF(tblDataIssues[狀態],"待確認")+COUNTIF(tblDataIssues[狀態],"待補資料")+COUNTIF(tblDataIssues[狀態],"待盤點")+COUNTIF(tblDataIssues[狀態],"待審查")', None),
            ("H4:I4", "H5:I6", "低庫存／負庫存", '=COUNTIF(tblStock[狀態],"補貨提醒")+COUNTIF(tblStock[狀態],"負庫存")', None),
            ("K4:L4", "K5:L6", "逾期回訪", '=COUNTIFS(tblFollowup[狀態],"<>已完成",tblFollowup[到期日],"<"&TODAY(),tblFollowup[到期日],"<>")', None),
            ("B8:C8", "B9:C10", "有效LINE對話", '=COUNTIF(tblLeads[漏斗階段],"有效對話")+COUNTIF(tblLeads[漏斗階段],"已加LINE")', None),
            ("E8:F8", "E9:F10", "已預約", '=COUNTIF(tblLeads[漏斗階段],"已預約")', None),
            ("H8:I8", "H9:I10", "已成交／回購", '=COUNTIF(tblLeads[漏斗階段],"已成交")+COUNTIF(tblLeads[漏斗階段],"回購")', None),
            ("K8:L8", "K9:L10", "累積實收", '=SUM(tblSales[實收金額])', self.money_fmt),
        ]
        for label_range, value_range, label, formula, numfmt in blocks:
            ws.merge_range(label_range, label, self.kpi_label)
            fmt = self.kpi_value
            if numfmt:
                fmt = self.wb.add_format(dict(fmt.properties)) if hasattr(fmt, "properties") else self.kpi_value
            first = value_range.split(":")[0]
            ws.merge_range(value_range, "", self.kpi_value)
            ws.write_formula(first, formula, numfmt or self.kpi_value, 0)
        ws.merge_range("B13:L13", "雙軌顧客旅程：每個入口都回到同一份CRM與回訪系統", self.section)
        professional = ["生活痛點內容", "加入LINE", "腦意識觀察", "深度諮詢", "課程／工具", "回訪與故事"]
        lifestyle = ["Moment故事／儀式", "加入LINE", "情境選擇", "蠟燭／首購組", "七日回訪", "檢測／回購／分享"]
        ws.write("B15", "專業路徑", self.kpi_label)
        ws.write("B17", "生活路徑", self.kpi_label)
        for idx, label in enumerate(professional, start=2):
            ws.write(14, idx, label, self.formula_fmt)
            if idx < 7:
                ws.write(14, idx + 1, "→", self.body)
        for idx, label in enumerate(lifestyle, start=2):
            ws.write(16, idx, label, self.input_fmt)
            if idx < 7:
                ws.write(16, idx + 1, "→", self.body)
        ws.merge_range("B20:L20", "本週決策提醒", self.section)
        reminders = [
            ("產品", '=COUNTIF(tblProducts[內容核准],"待確認")+COUNTIF(tblProducts[內容核准],"待老師核准")+COUNTIF(tblProducts[內容核准],"待建立")+COUNTIF(tblProducts[內容核准],"待補資料")+COUNTIF(tblProducts[內容核准],"待安全改寫")', "尚未核准的產品內容"),
            ("故事", '=COUNTIF(tblStories[狀態],"待核准")', "等待核准的故事"),
            ("內容", '=COUNTIF(tblContent[狀態],"待核准")', "等待核准的內容"),
            ("漏斗", '=COUNTIFS(tblLeads[狀態],"進行中",tblLeads[下次行動日],"<="&TODAY())', "今天需跟進的名單"),
        ]
        for row, (group, formula, label) in enumerate(reminders, start=21):
            ws.write(row - 1, 1, group, self.kpi_label)
            ws.write_formula(row - 1, 2, formula, self.kpi_value, 0)
            ws.merge_range(row - 1, 3, row - 1, 6, label, self.body)

    def build_roadmap(self) -> None:
        ws = self.add_sheet("02_90天路徑", "三階段不是日曆裝飾：每一列都有交付物、驗收標準、依賴與責任。目標日期由啟動日排定。", GOLD)
        headers = ["階段", "週次", "日程", "主題", "任務", "交付物", "KPI／驗收標準", "老師投入", "品牌規劃者產出", "負責", "依賴", "狀態", "目標日期", "證據／連結"]
        data = [row + ["未開始", "", ""] for row in ROADMAP]
        self.add_table(ws, 3, 0, "tblRoadmap", headers, data, style="Table Style Medium 4")
        widths = [18, 8, 13, 14, 38, 24, 50, 28, 30, 12, 20, 12, 13, 30]
        for col, width in enumerate(widths):
            ws.set_column(col, col, width)
        ws.data_validation(4, 11, 3 + len(data), 11, {"validate": "list", "source": "=GENERAL_STATUS"})
        ws.data_validation(4, 9, 3 + len(data), 9, {"validate": "list", "source": "=OWNER_LIST"})
        ws.conditional_format(4, 11, 3 + len(data), 11, {"type": "text", "criteria": "containing", "value": "已完成", "format": self.wb.add_format({"bg_color": PALE_GREEN, "font_color": GREEN})})
        ws.conditional_format(4, 11, 3 + len(data), 11, {"type": "text", "criteria": "containing", "value": "待", "format": self.wb.add_format({"bg_color": PALE_YELLOW})})
        ws.set_row(1, 36)

    def build_products(self) -> None:
        ws = self.add_sheet("03_產品主檔", "唯一正式產品來源。價格、成本與名稱未核准時保持『待確認』；對外報價只取老師核准列。", GOLD)
        headers = ["SKU", "產品類別", "正式名稱", "別名", "商業角色", "銷售狀態", "建議售價", "批發價", "單位成本", "單件毛利", "毛利率", "單位", "庫存管理", "補貨點", "來源依據", "內容核准", "備註"]
        data = [row[:9] + ["", ""] + row[9:] for row in PRODUCTS]
        formulas = {
            "單件毛利": '=IF(OR([@[建議售價]]="",[@[單位成本]]=""),"",[@[建議售價]]-[@[單位成本]])',
            "毛利率": '=IFERROR([@[單件毛利]]/[@[建議售價]],"")',
        }
        self.add_table(ws, 3, 0, "tblProducts", headers, data, formulas)
        widths = [18, 13, 35, 20, 22, 12, 12, 12, 12, 12, 10, 8, 10, 10, 32, 14, 45]
        for col, width in enumerate(widths):
            ws.set_column(col, col, width)
        for col in (6, 7, 8, 9):
            ws.set_column(col, col, widths[col], self.money_fmt)
        ws.set_column(10, 10, 10, self.percent_fmt)
        ws.data_validation(4, 12, 3 + len(data), 12, {"validate": "list", "source": "=YES_NO_UNKNOWN"})
        ws.conditional_format(4, 15, 3 + len(data), 15, {"type": "text", "criteria": "containing", "value": "待", "format": self.wb.add_format({"bg_color": PALE_YELLOW})})
        ws.conditional_format(4, 5, 3 + len(data), 5, {"type": "text", "criteria": "containing", "value": "待確認", "format": self.wb.add_format({"bg_color": PALE_RED, "font_color": RED})})

    def build_product_copy(self) -> None:
        ws = self.add_sheet("04_產品介紹", "內容規劃用，不是已核准商品頁。先核對事實、機制、規格與風險，再翻譯英文／日文。", NAVY_2)
        headers = ["SKU", "對外名稱", "客戶得到什麼", "使用情境", "獨特機制／事實邊界", "安全版一句話", "CTA", "仍需的證據", "規格候選", "宣稱風險", "審核狀態"]
        self.add_table(ws, 3, 0, "tblProductCopy", headers, PRODUCT_COPY, style="Table Style Medium 4")
        widths = [18, 30, 42, 34, 55, 48, 30, 55, 52, 10, 14]
        for col, width in enumerate(widths):
            ws.set_column(col, col, width)
        ws.data_validation(4, 9, 3 + len(PRODUCT_COPY), 9, {"validate": "list", "source": "=RISK_LEVEL"})
        ws.conditional_format(4, 9, 3 + len(PRODUCT_COPY), 9, {"type": "text", "criteria": "containing", "value": "高", "format": self.wb.add_format({"bg_color": PALE_RED, "font_color": RED, "bold": True})})
        ws.conditional_format(4, 10, 3 + len(PRODUCT_COPY), 10, {"type": "text", "criteria": "containing", "value": "待", "format": self.wb.add_format({"bg_color": PALE_YELLOW})})

    def build_inventory(self) -> None:
        ws = self.add_sheet("05_庫存異動", "只記一筆一筆的異動；數量一律填正數，由異動類型決定加減。實體盤點先建立『期初』。", NAVY_2)
        headers = ["異動ID", "日期", "SKU", "產品名稱", "異動類型", "數量", "單位成本", "金額", "關聯單號", "對象／供應商", "經手人", "備註"]
        rows = pad_rows([], 300, len(headers))
        formulas = {
            "產品名稱": '=IFERROR(INDEX(tblProducts[正式名稱],MATCH([@SKU],tblProducts[SKU],0)),"")',
            "單位成本": '=IFERROR(INDEX(tblProducts[單位成本],MATCH([@SKU],tblProducts[SKU],0)),"")',
            "金額": '=IF(OR([@數量]="",[@[單位成本]]=""),"",[@數量]*[@[單位成本]])',
        }
        self.add_table(ws, 3, 0, "tblInventory", headers, rows, formulas)
        widths = [16, 13, 18, 34, 14, 10, 12, 14, 18, 24, 14, 45]
        for col, width in enumerate(widths):
            ws.set_column(col, col, width)
        ws.set_column(1, 1, 13, self.date_fmt)
        ws.set_column(6, 7, 14, self.money_fmt)
        ws.data_validation(4, 2, 303, 2, {"validate": "list", "source": "=SKU_LIST"})
        ws.data_validation(4, 4, 303, 4, {"validate": "list", "source": "=INV_TYPE_LIST"})
        ws.data_validation(4, 10, 303, 10, {"validate": "list", "source": "=OWNER_LIST"})
        ws.conditional_format(4, 5, 303, 5, {"type": "cell", "criteria": "<=", "value": 0, "format": self.wb.add_format({"bg_color": PALE_RED})})

    def build_stock(self) -> None:
        ws = self.add_sheet("06_庫存總覽", "由05_庫存異動自動彙總。若顯示『待建立期初』，代表不是0庫存，而是尚未完成真實盤點。", NAVY_2)
        headers = ["SKU", "產品名稱", "單位成本", "補貨點", "期初", "進貨", "退貨入庫", "盤盈", "銷售", "退貨出庫", "贈品／樣品", "盤虧", "即時庫存", "庫存成本", "狀態", "最後異動"]
        physical = [row for row in PRODUCTS if row[10] == "是"]
        data = [[row[0], row[2], row[8], row[11], "", "", "", "", "", "", "", "", "", "", "", ""] for row in physical]
        formulas = {
            "期初": '=SUMIFS(tblInventory[數量],tblInventory[SKU],[@SKU],tblInventory[異動類型],"期初")',
            "進貨": '=SUMIFS(tblInventory[數量],tblInventory[SKU],[@SKU],tblInventory[異動類型],"進貨")',
            "退貨入庫": '=SUMIFS(tblInventory[數量],tblInventory[SKU],[@SKU],tblInventory[異動類型],"退貨入庫")',
            "盤盈": '=SUMIFS(tblInventory[數量],tblInventory[SKU],[@SKU],tblInventory[異動類型],"盤盈")',
            "銷售": '=SUMIFS(tblInventory[數量],tblInventory[SKU],[@SKU],tblInventory[異動類型],"銷售")',
            "退貨出庫": '=SUMIFS(tblInventory[數量],tblInventory[SKU],[@SKU],tblInventory[異動類型],"退貨出庫")',
            "贈品／樣品": '=SUMIFS(tblInventory[數量],tblInventory[SKU],[@SKU],tblInventory[異動類型],"贈品")+SUMIFS(tblInventory[數量],tblInventory[SKU],[@SKU],tblInventory[異動類型],"樣品")',
            "盤虧": '=SUMIFS(tblInventory[數量],tblInventory[SKU],[@SKU],tblInventory[異動類型],"盤虧")',
            "即時庫存": '=[@期初]+[@進貨]+[@退貨入庫]+[@盤盈]-[@銷售]-[@退貨出庫]-[@[贈品／樣品]]-[@盤虧]',
            "庫存成本": '=IF([@[單位成本]]="","",[@[即時庫存]]*[@[單位成本]])',
            "狀態": '=IF(COUNTIF(tblInventory[SKU],[@SKU])=0,"待建立期初",IF([@[即時庫存]]<0,"負庫存",IF([@[即時庫存]]<=[@補貨點],"補貨提醒","正常")))',
            "最後異動": '=IFERROR(MAXIFS(tblInventory[日期],tblInventory[SKU],[@SKU]),"")',
        }
        self.add_table(ws, 3, 0, "tblStock", headers, data, formulas, style="Table Style Medium 4")
        widths = [18, 34, 12, 10, 10, 10, 12, 10, 10, 12, 13, 10, 12, 14, 15, 13]
        for col, width in enumerate(widths):
            ws.set_column(col, col, width)
        ws.set_column(2, 2, 12, self.money_fmt)
        ws.set_column(13, 13, 14, self.money_fmt)
        ws.set_column(15, 15, 13, self.date_fmt)
        ws.conditional_format(4, 14, 3 + len(data), 14, {"type": "text", "criteria": "containing", "value": "負庫存", "format": self.wb.add_format({"bg_color": PALE_RED, "font_color": RED, "bold": True})})
        ws.conditional_format(4, 14, 3 + len(data), 14, {"type": "text", "criteria": "containing", "value": "補貨", "format": self.wb.add_format({"bg_color": PALE_YELLOW})})
        ws.conditional_format(4, 14, 3 + len(data), 14, {"type": "text", "criteria": "containing", "value": "待建立", "format": self.wb.add_format({"bg_color": PALE_BLUE})})

    def build_sales(self) -> None:
        ws = self.add_sheet("07_銷售訂單", "一列是一個訂單品項；同一訂單多品項可重複訂單ID。實體商品成交後，仍需在05_庫存異動記一筆『銷售』。", NAVY_2)
        headers = ["訂單ID", "訂單日期", "客戶ID", "客戶名稱", "通路", "來源內容ID", "SKU", "產品名稱", "數量", "單價", "折扣", "運費", "應收金額", "實收金額", "付款狀態", "履約狀態", "物流單號", "下次回訪日", "負責人", "備註"]
        rows = pad_rows([], 300, len(headers))
        formulas = {
            "客戶名稱": '=IFERROR(INDEX(tblCRM[姓名／稱呼],MATCH([@[客戶ID]],tblCRM[客戶ID],0)),"")',
            "產品名稱": '=IFERROR(INDEX(tblProducts[正式名稱],MATCH([@SKU],tblProducts[SKU],0)),"")',
            "單價": '=IFERROR(INDEX(tblProducts[建議售價],MATCH([@SKU],tblProducts[SKU],0)),"")',
            "應收金額": '=IF(OR([@數量]="",[@單價]=""),"",[@數量]*[@單價]-N([@折扣])+N([@運費]))',
        }
        self.add_table(ws, 3, 0, "tblSales", headers, rows, formulas)
        widths = [17, 13, 15, 18, 14, 16, 18, 34, 9, 12, 10, 10, 14, 14, 13, 13, 18, 13, 12, 40]
        for col, width in enumerate(widths):
            ws.set_column(col, col, width)
        ws.set_column(1, 1, 13, self.date_fmt)
        ws.set_column(9, 13, 14, self.money_fmt)
        ws.set_column(17, 17, 13, self.date_fmt)
        ws.data_validation(4, 4, 303, 4, {"validate": "list", "source": "=CHANNEL_LIST"})
        ws.data_validation(4, 6, 303, 6, {"validate": "list", "source": "=SKU_LIST"})
        ws.data_validation(4, 14, 303, 14, {"validate": "list", "source": "=PAYMENT_STATUS"})
        ws.data_validation(4, 15, 303, 15, {"validate": "list", "source": "=FULFILL_STATUS"})
        ws.data_validation(4, 18, 303, 18, {"validate": "list", "source": "=OWNER_LIST"})

    def build_crm(self) -> None:
        ws = self.add_sheet("08_客戶CRM", "只蒐集服務與回訪需要的資料。MEGAWING舊國外客戶未匯入；故事公開與行銷同意分開記錄。", NAVY_2)
        headers = ["客戶ID", "姓名／稱呼", "客戶類型", "狀態", "電話", "LINE", "Email", "IG", "縣市／國家", "生日", "性別", "年齡區間", "首次來源內容ID", "首次來源", "首次接觸日", "主要困擾", "事財情家體", "需求標籤", "身份／專長", "曾購產品", "最近服務", "最近聯絡日", "下次回訪日", "累積實收", "行銷同意", "故事公開同意", "負責人", "備註"]
        data = pad_rows(CRM_SEED, 300, len(headers))
        formulas = {"累積實收": '=IF([@[客戶ID]]="","",SUMIFS(tblSales[實收金額],tblSales[客戶ID],[@[客戶ID]]))'}
        self.add_table(ws, 3, 0, "tblCRM", headers, data, formulas, style="Table Style Medium 4")
        widths = [15, 18, 18, 13, 15, 16, 25, 18, 16, 12, 9, 12, 18, 18, 13, 30, 12, 34, 22, 30, 26, 13, 13, 14, 12, 16, 12, 45]
        for col, width in enumerate(widths):
            ws.set_column(col, col, width)
        for col in (9, 14, 21, 22):
            ws.set_column(col, col, widths[col], self.date_fmt)
        ws.set_column(23, 23, 14, self.money_fmt)
        ws.data_validation(4, 2, 303, 2, {"validate": "list", "source": "=CUSTOMER_TYPE"})
        ws.data_validation(4, 3, 303, 3, {"validate": "list", "source": "=CRM_STATUS"})
        ws.data_validation(4, 13, 303, 13, {"validate": "list", "source": "=CHANNEL_LIST"})
        ws.data_validation(4, 16, 303, 16, {"validate": "list", "source": "=FIVE_DIM_LIST"})
        ws.data_validation(4, 24, 303, 25, {"validate": "list", "source": "=YES_NO_UNKNOWN"})
        ws.data_validation(4, 26, 303, 26, {"validate": "list", "source": "=OWNER_LIST"})
        ws.conditional_format(4, 22, 303, 22, {"type": "formula", "criteria": '=AND($W5<TODAY(),$W5<>"")', "format": self.wb.add_format({"bg_color": PALE_YELLOW})})

    def build_service(self) -> None:
        ws = self.add_sheet("09_服務紀錄", "一列是一回服務或互動摘要。記錄生活議題與客戶自述，不在表內自行下醫療診斷。", NAVY_2)
        headers = ["紀錄ID", "日期", "客戶ID", "客戶名稱", "服務SKU", "服務名稱", "主要議題", "事財情家體", "服務摘要", "觀察資料／檔案", "建議下一步", "推薦SKU", "回訪日", "客戶自述觀察", "隱私等級", "負責人", "備註"]
        rows = pad_rows([], 300, len(headers))
        formulas = {
            "客戶名稱": '=IFERROR(INDEX(tblCRM[姓名／稱呼],MATCH([@[客戶ID]],tblCRM[客戶ID],0)),"")',
            "服務名稱": '=IFERROR(INDEX(tblProducts[正式名稱],MATCH([@[服務SKU]],tblProducts[SKU],0)),"")',
        }
        self.add_table(ws, 3, 0, "tblService", headers, rows, formulas)
        widths = [16, 13, 15, 18, 18, 32, 30, 12, 45, 35, 35, 18, 13, 45, 12, 12, 40]
        for col, width in enumerate(widths):
            ws.set_column(col, col, width)
        ws.set_column(1, 1, 13, self.date_fmt)
        ws.set_column(12, 12, 13, self.date_fmt)
        ws.data_validation(4, 4, 303, 4, {"validate": "list", "source": "=SKU_LIST"})
        ws.data_validation(4, 7, 303, 7, {"validate": "list", "source": "=FIVE_DIM_LIST"})
        ws.data_validation(4, 11, 303, 11, {"validate": "list", "source": "=SKU_LIST"})
        ws.data_validation(4, 14, 303, 14, {"validate": "list", "source": "=PRIVACY_LEVEL"})
        ws.data_validation(4, 15, 303, 15, {"validate": "list", "source": "=OWNER_LIST"})

    def build_followup(self) -> None:
        ws = self.add_sheet("10_回訪任務", "把『記得再聯絡』變成有日期、有腳本、有結果的任務；完成後可建立下一次任務。", NAVY_2)
        headers = ["任務ID", "客戶ID", "客戶名稱", "觸發類型", "來源紀錄ID", "到期日", "任務內容", "建議訊息", "狀態", "完成日", "結果", "下一任務日", "負責人", "備註"]
        rows = pad_rows([], 300, len(headers))
        formulas = {"客戶名稱": '=IFERROR(INDEX(tblCRM[姓名／稱呼],MATCH([@[客戶ID]],tblCRM[客戶ID],0)),"")'}
        self.add_table(ws, 3, 0, "tblFollowup", headers, rows, formulas, style="Table Style Medium 4")
        widths = [16, 15, 18, 18, 18, 13, 34, 55, 13, 13, 30, 13, 12, 40]
        for col, width in enumerate(widths):
            ws.set_column(col, col, width)
        for col in (5, 9, 11):
            ws.set_column(col, col, 13, self.date_fmt)
        ws.data_validation(4, 8, 303, 8, {"validate": "list", "source": "=GENERAL_STATUS"})
        ws.data_validation(4, 12, 303, 12, {"validate": "list", "source": "=OWNER_LIST"})
        ws.conditional_format(4, 5, 303, 5, {"type": "formula", "criteria": '=AND($F5<TODAY(),$I5<>"已完成",$F5<>"")', "format": self.wb.add_format({"bg_color": PALE_RED, "font_color": RED})})

    def build_content(self) -> None:
        ws = self.add_sheet("11_內容主題庫", "40個種子題目都有任務、產品角色、CTA、證據與風險；故事題只能換入真實案例，不能照題目虛構。", NAVY_2)
        headers = ["內容ID", "帳號", "漏斗階段", "內容類型", "標題／Hook", "客戶議題", "產品角色", "SKU", "格式", "CTA", "誘因／承接", "所需證據", "宣稱風險", "狀態", "預定拍攝日", "預定發布日", "實際發布日", "連結", "負責人", "備註"]
        data = [row + ["想法", "", "", "", "", "", ""] for row in CONTENT_IDEAS]
        self.add_table(ws, 3, 0, "tblContent", headers, data)
        widths = [14, 18, 12, 13, 52, 25, 18, 18, 12, 28, 28, 45, 10, 13, 13, 13, 13, 30, 12, 40]
        for col, width in enumerate(widths):
            ws.set_column(col, col, width)
        ws.data_validation(4, 3, 3 + len(data), 3, {"validate": "list", "source": "=CONTENT_TYPE"})
        ws.data_validation(4, 7, 3 + len(data), 7, {"validate": "list", "source": "=SKU_LIST"})
        ws.data_validation(4, 12, 3 + len(data), 12, {"validate": "list", "source": "=RISK_LEVEL"})
        ws.data_validation(4, 13, 3 + len(data), 13, {"validate": "list", "source": "=CONTENT_STATUS"})
        ws.data_validation(4, 18, 3 + len(data), 18, {"validate": "list", "source": "=OWNER_LIST"})
        ws.conditional_format(4, 12, 3 + len(data), 12, {"type": "text", "criteria": "containing", "value": "高", "format": self.wb.add_format({"bg_color": PALE_RED, "font_color": RED})})

    def build_content_metrics(self) -> None:
        ws = self.add_sheet("12_內容成效", "觀看不是終點。每支內容用同一內容ID追到LINE新增、預約、訂單與實收，才能判斷是否值得放大。", NAVY_2)
        headers = ["內容ID", "發布日", "平台", "觀看", "3秒觀看", "平均觀看秒", "完播率", "收藏", "分享", "留言", "私訊", "LINE新增", "預約", "訂單", "實收", "廣告費", "互動率", "LINE轉換率", "預約轉換率", "成交轉換率", "備註"]
        rows = pad_rows([], 200, len(headers))
        formulas = {
            "互動率": '=IFERROR(([@收藏]+[@分享]+[@留言]+[@私訊]) / [@觀看],"")',
            "LINE轉換率": '=IFERROR([@LINE新增]/[@觀看],"")',
            "預約轉換率": '=IFERROR([@預約]/[@LINE新增],"")',
            "成交轉換率": '=IFERROR([@訂單]/[@LINE新增],"")',
        }
        self.add_table(ws, 3, 0, "tblContentMetrics", headers, rows, formulas, style="Table Style Medium 4")
        widths = [14, 13, 14, 11, 11, 13, 11, 10, 10, 10, 10, 11, 10, 10, 14, 12, 12, 13, 14, 14, 40]
        for col, width in enumerate(widths):
            ws.set_column(col, col, width)
        ws.set_column(1, 1, 13, self.date_fmt)
        ws.set_column(14, 15, 14, self.money_fmt)
        for col in range(16, 20):
            ws.set_column(col, col, widths[col], self.percent_fmt)
        ws.data_validation(4, 2, 203, 2, {"validate": "list", "source": "=PLATFORM_LIST"})

    def build_stories(self) -> None:
        ws = self.add_sheet("13_希望故事庫", "每則故事都必須可回溯來源、取得同意、完成匿名與風險審查。客戶主觀感受不等於產品功效證明。", NAVY_2)
        headers = ["故事ID", "客戶ID", "匿名稱呼", "來源日期", "來源管道", "產品／服務", "原先情境", "採取的行動", "客戶自述觀察", "原始證據路徑", "內部使用同意", "匿名公開同意", "已去識別", "核准文案", "審核人", "風險", "狀態", "發布日", "發布連結", "撤回日", "備註"]
        rows = pad_rows([], 150, len(headers))
        self.add_table(ws, 3, 0, "tblStories", headers, rows)
        widths = [14, 15, 18, 13, 15, 28, 42, 38, 45, 40, 15, 16, 13, 60, 12, 10, 13, 13, 30, 13, 40]
        for col, width in enumerate(widths):
            ws.set_column(col, col, width)
        ws.data_validation(4, 10, 153, 12, {"validate": "list", "source": "=YES_NO_UNKNOWN"})
        ws.data_validation(4, 15, 153, 15, {"validate": "list", "source": "=RISK_LEVEL"})
        ws.data_validation(4, 16, 153, 16, {"validate": "list", "source": "=GENERAL_STATUS"})
        ws.conditional_format(4, 15, 153, 15, {"type": "text", "criteria": "containing", "value": "高", "format": self.wb.add_format({"bg_color": PALE_RED, "font_color": RED})})

    def build_leads(self) -> None:
        ws = self.add_sheet("14_LINE漏斗", "一列是一個名單／對話。記錄來源內容ID與下一步，才能知道短影音是否真正帶來有效對話與成交。", NAVY_2)
        headers = ["名單ID", "進線日", "客戶ID", "來源內容ID", "通路", "關鍵字", "主要需求", "漏斗階段", "首次回覆時間", "回覆分鐘", "下一步", "下次行動日", "預約日", "訂單ID", "流失原因", "負責人", "狀態", "備註"]
        rows = pad_rows([], 300, len(headers))
        self.add_table(ws, 3, 0, "tblLeads", headers, rows)
        widths = [15, 13, 15, 16, 14, 15, 30, 13, 18, 12, 30, 13, 13, 17, 30, 12, 13, 40]
        for col, width in enumerate(widths):
            ws.set_column(col, col, width)
        ws.data_validation(4, 4, 303, 4, {"validate": "list", "source": "=CHANNEL_LIST"})
        ws.data_validation(4, 7, 303, 7, {"validate": "list", "source": "=FUNNEL_STAGE"})
        ws.data_validation(4, 15, 303, 15, {"validate": "list", "source": "=OWNER_LIST"})
        ws.data_validation(4, 16, 303, 16, {"validate": "list", "source": "=GENERAL_STATUS"})

    def build_line_sop(self) -> None:
        ws = self.add_sheet("15_LINE回覆SOP", "先用情境與資格提問理解需求，再給一個下一步。高風險健康需求要升級或轉介，不能用產品承諾承接。", NAVY_2)
        headers = ["情境", "觸發／關鍵字", "目的", "第一回覆", "資格提問", "下一個CTA", "LINE標籤", "回覆SLA", "禁止說法", "狀態"]
        self.add_table(ws, 3, 0, "tblLineSOP", headers, LINE_SOPS, style="Table Style Medium 4")
        widths = [20, 24, 32, 68, 48, 38, 20, 16, 45, 14]
        for col, width in enumerate(widths):
            ws.set_column(col, col, width)
        ws.conditional_format(4, 8, 3 + len(LINE_SOPS), 8, {"type": "text", "criteria": "containing", "value": "治療", "format": self.wb.add_format({"bg_color": PALE_RED})})

    def build_profit_share(self) -> None:
        ws = self.add_sheet("16_分潤試算", "這是決策試算，不是已成立契約。先填結算期間、基準營收、費率與實際成本，再以書面合作條款為準。", GOLD)
        ws.freeze_panes(13, 0)
        ws.set_column("A:A", 3)
        ws.set_column("B:B", 24)
        ws.set_column("C:C", 18)
        ws.set_column("D:F", 16)
        ws.merge_range("B4:F4", "結算設定（黃色欄由雙方確認）", self.section)
        settings = [
            ("結算開始日", "", "填日期"),
            ("結算結束日", "", "填日期"),
            ("固定顧問費", 0, "不與營收保證綁定"),
            ("基準營收", 0, "績效獎金只計超過基準的部分"),
            ("績效獎金率", 0, "例如5%填5%"),
            ("預估其他費用", 0, "稅、金流、廣告或經雙方同意費用"),
        ]
        for row, (label, value, note) in enumerate(settings, start=5):
            ws.write(row - 1, 1, label, self.body)
            ws.write(row - 1, 2, value, self.input_fmt)
            ws.merge_range(row - 1, 3, row - 1, 5, note, self.body)
        ws.set_column("C:C", 18, self.money_fmt)
        ws.write("C9", 0, self.percent_fmt)
        ws.merge_range("H4:J4", "結算摘要（公式）", self.section)
        summary = [
            ("期間實收", '=SUM(tblShare[期間實收])', self.money_fmt),
            ("產品成本", '=SUM(tblShare[產品成本])', self.money_fmt),
            ("期間毛利", '=SUM(tblShare[期間毛利])', self.money_fmt),
            ("超額營收", '=MAX(0,I5-$C$8)', self.money_fmt),
            ("績效獎金", '=I8*$C$9', self.money_fmt),
            ("新產品分潤", '=SUM(tblShare[新產品分潤])', self.money_fmt),
            ("應付規劃者", '=$C$7+I9+I10', self.money_fmt),
        ]
        for row, (label, formula, fmt) in enumerate(summary, start=5):
            ws.write(row - 1, 7, label, self.kpi_label)
            ws.write_formula(row - 1, 8, formula, fmt, 0)
            ws.merge_range(row - 1, 8, row - 1, 9, "", fmt)
            ws.write_formula(row - 1, 8, formula, fmt, 0)
        headers = ["SKU", "產品名稱", "期間數量", "期間實收", "單位成本", "產品成本", "期間毛利", "是否新產品", "新產品分潤率", "新產品分潤", "備註"]
        data = [[row[0], row[2], "", "", row[8], "", "", "否", 0, "", "成本與新產品定義須書面確認"] for row in PRODUCTS]
        formulas = {
            "期間數量": '=IF(OR($C$5="",$C$6=""),0,SUMIFS(tblSales[數量],tblSales[SKU],[@SKU],tblSales[訂單日期],">="&$C$5,tblSales[訂單日期],"<="&$C$6,tblSales[付款狀態],"已付款"))',
            "期間實收": '=IF(OR($C$5="",$C$6=""),0,SUMIFS(tblSales[實收金額],tblSales[SKU],[@SKU],tblSales[訂單日期],">="&$C$5,tblSales[訂單日期],"<="&$C$6,tblSales[付款狀態],"已付款"))',
            "產品成本": '=IF([@[單位成本]]="",0,[@[期間數量]]*[@[單位成本]])',
            "期間毛利": '=[@[期間實收]]-[@[產品成本]]',
            "新產品分潤": '=IF([@[是否新產品]]="是",MAX(0,[@[期間毛利]])*[@[新產品分潤率]],0)',
        }
        self.add_table(ws, 12, 0, "tblShare", headers, data, formulas)
        widths = [18, 35, 12, 14, 12, 14, 14, 13, 15, 14, 45]
        for col, width in enumerate(widths):
            ws.set_column(col, col, width)
        for col in (3, 4, 5, 6, 9):
            ws.set_column(col, col, widths[col], self.money_fmt)
        ws.set_column(8, 8, 15, self.percent_fmt)
        ws.data_validation(13, 7, 12 + len(data), 7, {"validate": "list", "source": "=YES_NO_UNKNOWN"})

    def build_decisions(self) -> None:
        ws = self.add_sheet("17_品牌決策", "建議不等於老師已核准。所有會改變品牌方向、價格、宣稱、資料使用或合作權益的項目，都在這裡留下決策。", GOLD)
        headers = ["決策ID", "類別", "要決定什麼", "本版建議", "理由／依據", "最晚決策點", "決策者", "狀態", "最終決定", "決策日期", "證據／會議"]
        data = [row + ["", "", ""] for row in DECISIONS]
        self.add_table(ws, 3, 0, "tblDecisions", headers, data, style="Table Style Medium 4")
        widths = [14, 12, 30, 50, 50, 15, 12, 13, 55, 13, 30]
        for col, width in enumerate(widths):
            ws.set_column(col, col, width)
        ws.data_validation(4, 6, 3 + len(data), 6, {"validate": "list", "source": "=OWNER_LIST"})
        ws.data_validation(4, 7, 3 + len(data), 7, {"validate": "list", "source": "=GENERAL_STATUS"})
        ws.conditional_format(4, 7, 3 + len(data), 7, {"type": "text", "criteria": "containing", "value": "待", "format": self.wb.add_format({"bg_color": PALE_YELLOW})})

    def build_audit(self) -> None:
        ws = self.add_sheet("18_資料核對", "本頁保留不確定性與來源處理紀錄。『已隔離』代表未匯入正式系統，不代表原始檔被刪除。", RED)
        issue_headers = ["問題ID", "嚴重度", "類別", "發現", "影響", "處理建議", "來源", "狀態"]
        self.add_table(ws, 3, 0, "tblDataIssues", issue_headers, DATA_ISSUES, style="Table Style Medium 3")
        widths = [14, 10, 14, 55, 45, 55, 42, 15]
        for col, width in enumerate(widths):
            ws.set_column(col, col, width)
        ws.data_validation(4, 1, 3 + len(DATA_ISSUES), 1, {"validate": "list", "source": "=RISK_LEVEL"})
        ws.conditional_format(4, 1, 3 + len(DATA_ISSUES), 1, {"type": "text", "criteria": "containing", "value": "高", "format": self.wb.add_format({"bg_color": PALE_RED, "font_color": RED, "bold": True})})
        source_start = 7 + len(DATA_ISSUES)
        ws.merge_range(source_start - 1, 0, source_start - 1, 7, "來源檔採用／隔離紀錄", self.section)
        source_headers = ["來源檔", "用途", "處理", "說明"]
        self.add_table(ws, source_start, 0, "tblSourceAudit", source_headers, SOURCE_AUDIT, style="Table Style Medium 4")
        ws.set_column(0, 0, 52)
        ws.set_column(1, 1, 22)
        ws.set_column(2, 2, 15)
        ws.set_column(3, 3, 70)

    def close(self) -> None:
        guide = self.wb.get_worksheet_by_name("00_使用說明")
        lists = self.wb.get_worksheet_by_name("19_選單")
        if guide is not None:
            guide.activate()
            guide.select()
        if lists is not None:
            lists.hide()
        self.wb.close()


def build_workbook(path: Path) -> None:
    builder = WorkbookBuilder(path)
    builder.build_guide()
    builder.build_dashboard()
    builder.build_roadmap()
    builder.build_products()
    builder.build_product_copy()
    builder.build_inventory()
    builder.build_stock()
    builder.build_sales()
    builder.build_crm()
    builder.build_service()
    builder.build_followup()
    builder.build_content()
    builder.build_content_metrics()
    builder.build_stories()
    builder.build_leads()
    builder.build_line_sop()
    builder.build_profit_share()
    builder.build_decisions()
    builder.build_audit()
    builder.build_lists()
    builder.close()


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: object, bold: bool = False, color: str = TEXT, size: int = 9) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Microsoft JhengHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_doc_table(doc: Document, headers: list[str], rows: list[list[object]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for idx, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[idx], NAVY_2)
        set_cell_text(table.rows[0].cells[idx], header, True, WHITE, 9)
    for row_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for col_idx, value in enumerate(row):
            if row_idx % 2 == 1:
                set_cell_shading(cells[col_idx], LIGHT_GREY)
            set_cell_text(cells[col_idx], value, False, TEXT, 8)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def build_document(path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft JhengHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(6)
    for name, size, color in [("Title", 28, NAVY), ("Heading 1", 18, NAVY), ("Heading 2", 13, NAVY_2), ("Heading 3", 11, GOLD)]:
        style = styles[name]
        style.font.name = "Microsoft JhengHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True

    if "Callout" not in styles:
        callout = styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
    else:
        callout = styles["Callout"]
    callout.font.name = "Microsoft JhengHei"
    callout._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    callout.font.size = Pt(12)
    callout.font.bold = True
    callout.font.color.rgb = RGBColor.from_string(NAVY)
    callout.paragraph_format.space_before = Pt(8)
    callout.paragraph_format.space_after = Pt(8)
    callout.paragraph_format.left_indent = Cm(0.7)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(80)
    run = title.add_run("希望之光品牌規劃路徑圖")
    run.font.name = "Microsoft JhengHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(30)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(NAVY)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("從短影音委託，升級為可追蹤、可回購、可交接的品牌營運系統").bold = True
    doc.add_paragraph()
    call = doc.add_paragraph(style="Callout")
    call.alignment = WD_ALIGN_PARAGRAPH.CENTER
    call.add_run("看見自己的光，整理一條客戶走得懂、團隊接得住的路。")
    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("初版日期：2026-08-09\n").bold = True
    meta.add_run("整理：[Codex]｜依使用者提供之16份來源檔交叉核對\n")
    meta.add_run("重要：文中『建議』不是老師已核准決策；價格、庫存、規格與健康相關文案須完成待確認清單。")
    doc.add_page_break()

    doc.add_heading("先說結論：希望之光需要的不是更多影片，而是一條完整的成交與回流路徑", level=1)
    p = doc.add_paragraph(style="Callout")
    p.add_run("核心策略：老師是信任核心；希望之光是承接品牌；腦意識觀察與五分鐘儀式是兩個入口；LINE、CRM、服務紀錄與回訪是同一套後台。")
    add_bullets(
        doc,
        [
            "短影音只負責讓對的人停下來，不能單獨負責成交。每一支內容必須有一個明確下一步與可追蹤內容ID。",
            "陌生客的專業入口以腦意識觀察／檢測建立信任；生活型受眾可從蠟燭與五分鐘儀式進入。兩條路都回到官方LINE。",
            "大小貴人、諧和機與課程不是同層商品：它們要依客戶已建立的理解、需求與投入程度承接，而非同時出現在同一個銷售訊息裡。",
            "上千份手寫個案與既有訊息是資產，但只有在完成資料最小化、同意、匿名、來源與回訪日期後，才是可長期使用的CRM與故事庫。",
            "90天的成功先看系統是否建好、資料是否可追、流程是否被團隊實際使用；營收是重要結果，但不能由品牌規劃者單方面保證。",
        ],
    )

    doc.add_heading("一、現況診斷：四個真正的阻力", level=1)
    add_doc_table(
        doc,
        ["阻力", "來源顯示的現象", "商業後果", "本案處理"],
        [
            ["內容很多，路徑不明", "約37支影片多為老師講述／個案分析，觀眾不知道下一步", "觀看與成交斷裂", "每支內容配置單一CTA、LINE關鍵字與來源ID"],
            ["產品同時出現，角色混亂", "檢測、蠟燭、大小貴人、課程與機器沒有一致排序", "客戶不知道從哪裡開始", "建立雙入口與產品階梯"],
            ["黃金資產未活化", "上千份手寫個案、約800則訊息、故事與見證未結構化", "舊客無法回流、內容每次重做", "CRM、故事同意、回訪任務與批次數位化"],
            ["營運表格不可依賴", "舊模板污染、#REF!、價格互換、50%／60%成本混用", "庫存、毛利、分潤與目標失真", "重新建立SKU、異動式庫存與待確認機制"],
        ],
    )

    doc.add_heading("二、品牌核心與架構", level=1)
    doc.add_heading("2.1 建議的品牌一句話", level=2)
    p = doc.add_paragraph(style="Callout")
    p.add_run("陪你看見自己，找到適合自己的節奏。")
    doc.add_paragraph("這句話能容納腦意識觀察、人生梳理、課程、蠟燭與調頻工具，又不把品牌限制在單一產品或醫療效果。老師仍需在Day 5-7正式核准。")
    doc.add_heading("2.2 品牌架構", level=2)
    add_doc_table(
        doc,
        ["層級", "角色", "內容／產品", "主要任務"],
        [
            ["信任核心", "珈語老師本人", "觀點、經驗、提問與陪伴方式", "讓人相信這套看見與整理的方法"],
            ["母品牌", "希望之光 Hope Light", "統一承接所有服務、產品、故事與資料", "避免希望之光／希望之盒或多帳號各自為政"],
            ["專業引擎", "hopelight.ig", "腦意識觀察、個案知識、課程與老師世界觀", "負責『看見』與專業信任"],
            ["生活／信任引擎", "hopelight.moment", "五分鐘儀式、蠟燭、大小貴人、諧和機、真實故事", "負責『陪伴』、社會證明與低門檻互動"],
            ["後台系統", "LINE＋CRM＋訂單＋庫存＋回訪", "所有客戶與行為資料", "承接、追蹤、回購與交接"],
        ],
    )
    doc.add_heading("2.3 品牌語言規則", level=2)
    add_doc_table(
        doc,
        ["避免直接說", "改用生活語言", "理由"],
        [
            ["治療失眠／焦慮／憂鬱／ADHD", "建立睡前收尾、整理思緒、陪伴自我觀察；必要時優先尋求合格專業協助", "避免醫療診斷與療效承諾"],
            ["系統校準人生頻率", "當你更了解自己，做決定會更有方向", "把抽象專業翻譯成客戶能理解的改變"],
            ["用了就招財／招桃花／有效", "使用者分享她的感受與行動；產品是一個提醒自己的媒介", "保留故事真實性，不把主觀回饋當必然結果"],
            ["老師很厲害、技術很多", "你適不適合？你會經歷什麼？下一步怎麼開始？", "每個入口都回答客戶最在意的問題"],
        ],
    )

    doc.add_page_break()
    doc.add_heading("三、產品角色與雙軌顧客旅程", level=1)
    doc.add_heading("3.1 產品角色", level=2)
    add_doc_table(
        doc,
        ["角色", "產品／服務", "客戶此時需要", "下一步"],
        [
            ["專業入口", "腦意識觀察／檢測90分鐘｜候選價2,980", "第一次結構化看見自己的思維與生活議題", "深度諮詢或合適的日常陪伴"],
            ["生活入口", "五分鐘儀式內容、頻率蠟燭、21顆首購組", "用較低心理門檻開始一個可重複的生活儀式", "七日回訪、故事、檢測或回購"],
            ["建立信任", "60分鐘諮詢、大小貴人與其他調頻工具", "把已看見的問題整理得更具體，或延伸到日常", "課程、長期練習或合適工具"],
            ["核心服務", "2天／3天課程", "把一次觀察轉成較完整的理解與方法", "社群、複訓、轉介紹或持續陪伴"],
            ["延續陪伴", "蠟燭、大小貴人、諧和腦波訓練機", "在生活中持續保留整理與練習的提示", "回購、回訪、故事回流"],
        ],
    )
    doc.add_heading("3.2 專業路徑 A", level=2)
    doc.add_paragraph("生活痛點內容 → 加入LINE／回答三題 → 90分鐘腦意識觀察 → 個人化整理 → 60分鐘諮詢或合適課程／工具 → 7／30日回訪 → 真實故事或轉介紹")
    doc.add_heading("3.3 生活路徑 B", level=2)
    doc.add_paragraph("Hope Light Moment故事／儀式 → 留言關鍵字 → LINE情境選擇 → 蠟燭或首購組 → 安全使用引導 → 7日回訪 → 回購／檢測／故事分享")
    doc.add_heading("3.4 每一步的成功條件", level=2)
    add_doc_table(
        doc,
        ["節點", "一定要有", "追蹤指標", "常見斷點"],
        [
            ["內容", "一個受眾、一個問題、一個CTA、一個內容ID", "停留、收藏、私訊、LINE新增", "同一支片同時推多個產品"],
            ["LINE", "關鍵字、第一回覆、資格提問、下一步與負責人", "回覆時間、有效對話率", "回完資訊但沒有下一個行動"],
            ["預約／購買", "正式價格、包含內容、界線、付款與取消方式", "預約率、付款率、流失原因", "價格不一致或用效果承諾成交"],
            ["交付", "服務／出貨紀錄與下一次回訪日", "完成率、客訴、退換", "做完服務就失聯"],
            ["回訪", "固定7／30日節奏與非誘導式問題", "回購率、轉介紹、可用故事", "只問『有沒有效』"],
        ],
    )

    doc.add_heading("四、90天路徑圖", level=1)
    add_doc_table(
        doc,
        ["階段", "目的", "核心成果", "過關條件"],
        [
            ["Day 1-30｜梳理", "把混亂變成單一版本與清楚路徑", "定位、產品主檔、雙軌旅程、CRM、LINE v1、40題內容庫、期初盤點", "價格與名稱矛盾已決策；所有實體SKU有期初；兩條路可完整走測"],
            ["Day 31-60｜驗證", "用小批次證明哪些內容、客群與產品路徑有效", "兩輪內容、兩批舊客回流、產品實驗、七日回訪、週報", "能把內容ID追到LINE、預約、訂單；停止無效假設"],
            ["Day 61-90｜放大", "把有效做法寫成團隊能使用的SOP", "內容、客服、訂單庫存、故事、分潤、儀表板與交接", "團隊可不靠規劃者完成一次循環；91-180天只放大已驗證項"],
        ],
    )
    doc.add_heading("4.1 前14天，現在就照這個順序做", level=2)
    add_bullets(
        doc,
        [
            "Day 1：與老師確認90天最重要的單一商業問題，以及誰有最終核准權。",
            "Day 1-2：把所有來源列為採用、隔離或待確認；本次已完成第一版。",
            "Day 2-5：訪談品牌起點、真實故事、產品獨特流程、理想客戶與半年理想畫面。",
            "Day 3-5：取得IG、LINE、近3個月銷售、目前庫存與訊息基準，沒有資料就明確標缺口。",
            "Day 5-7：老師核准母品牌、品牌一句話、醫療／效果宣稱界線。",
            "Day 8-10：逐SKU確認名稱、售價、成本、規格、是否在售；60分鐘諮詢已確認為1,980，接著關閉大小貴人價格矛盾。",
            "Day 10-12：定稿專業路徑與生活路徑；每一步指定CTA、承接人、資料欄與完成條件。",
            "Day 12-14：上線Bio／連結架構，開始CRM試轉與內容拍攝準備；此時才開始正式量產短影音。",
        ],
    )
    doc.add_heading("4.2 每週會議只回答五題", level=2)
    add_bullets(
        doc,
        [
            "本週哪一個內容帶來最多有效對話，而不是最多空泛觀看？",
            "哪個環節掉得最多：看到、加入LINE、回覆、預約、付款、回訪？",
            "哪個客群最容易往下一步走？原因有沒有實際紀錄？",
            "哪個產品有毛利、有庫存、有人承接，也有後續陪伴？",
            "下週只放大什麼、停止什麼、還缺哪一個決策？",
        ],
    )

    doc.add_page_break()
    doc.add_heading("五、資料與表格系統：讓每個動作可追蹤", level=1)
    add_doc_table(
        doc,
        ["資料層", "唯一主檔／交易檔", "關鍵規則"],
        [
            ["產品", "03_產品主檔、04_產品介紹", "正式名稱、價格、成本只維護一次；內容核准與事實資料分開"],
            ["庫存", "05_庫存異動、06_庫存總覽", "不再橫向新增日期欄；所有數量用異動類型計算"],
            ["銷售", "07_銷售訂單", "每列一個訂單品項；同一訂單可多列；實收與付款狀態分開"],
            ["客戶", "08_客戶CRM", "一人一ID；只蒐集服務所需資料；同意狀態可追溯"],
            ["服務與回訪", "09_服務紀錄、10_回訪任務", "每次服務留下下一步；回訪一定有日期、結果與下一任務"],
            ["內容與故事", "11-13工作表", "內容ID追到轉換；故事先留原始證據，再匿名、同意、核准"],
            ["LINE漏斗", "14_LINE漏斗、15_LINE回覆SOP", "每個對話有來源、階段、下一步與負責人"],
            ["商業決策", "16_分潤試算、17_品牌決策、18_資料核對", "未確認成本不結算；建議與正式決定分開"],
        ],
    )
    doc.add_heading("六、內容與證據系統", level=1)
    add_doc_table(
        doc,
        ["內容類型", "建議比例", "任務", "主要CTA"],
        [
            ["真實故事", "35%", "讓人看見生活中的真實轉變與陪伴方式", "留言情境字／加入LINE"],
            ["常見問題", "25%", "降低第一次接觸、使用與價格的不確定", "查看流程／取得指南"],
            ["知識", "20%", "把專業翻成生活語言，不堆砌術語", "收藏／加入LINE"],
            ["幕後", "10%", "讓獨特機制與用心可被看見", "了解產品流程"],
            ["見證", "10%", "累積社會證明，但保留來源、同意與主觀語氣", "查看故事庫／分享自己的Moment"],
        ],
    )
    doc.add_paragraph("本案已在Excel中放入40個內容題目。所有『今天的希望故事』都只是採訪模板，必須換入真實案例、取得同意並由老師核准，不能直接當成已發生故事。")

    doc.add_heading("七、KPI：先量系統，再量轉換，最後量收入", level=1)
    add_doc_table(
        doc,
        ["層級", "指標", "定義／算法", "90天用途"],
        [
            ["系統建置", "產品主檔核准率", "已核准在售SKU ÷ 全部在售SKU", "Day 30前消除報價與名稱混亂"],
            ["資料資產", "名單數位化完成率", "已人工校正CRM筆數 ÷ 本期目標筆數", "先用30-50份試點，再擴大"],
            ["內容", "有效對話率", "由內容帶來的有效LINE對話 ÷ 觀看或觸及", "判斷Hook與CTA是否吸引對的人"],
            ["漏斗", "LINE→預約率", "預約人數 ÷ 有效LINE對話數", "找出承接與服務說明斷點"],
            ["成交", "預約→付款率", "已付款訂單 ÷ 有效預約", "檢查價格、信任與適配"],
            ["留存", "7日回訪完成率", "到期且完成回訪 ÷ 到期回訪", "確保售後不是口頭承諾"],
            ["回購", "90日回購率", "期間內再次購買客戶 ÷ 已購客戶", "衡量陪伴與產品週期"],
            ["商業", "毛利與可歸因收入", "實收－單位成本－明確變動費；需能對回來源", "分潤與放大前的基礎"],
        ],
    )

    doc.add_heading("八、角色分工與合作邊界", level=1)
    add_doc_table(
        doc,
        ["事項", "老師", "品牌規劃者", "共同完成"],
        [
            ["品牌與產品事實", "提供真實說法、證據並做最後核准", "訪談、整理、翻譯與提出選項", "形成可對外的一致版本"],
            ["內容", "出鏡、提供故事與專業判斷", "策略、腳本、拍攝編排、CTA與數據", "每週決定保留與停止"],
            ["客戶服務", "處理專業判斷與高風險情境", "建立SOP、欄位與追蹤", "依同一套界線回覆"],
            ["銷售與庫存", "確認價格、成本、供貨與實物盤點", "建立表格、流程、毛利與異常檢查", "按月對帳"],
            ["成效", "投入服務與及時核准", "提供可追蹤系統與分析", "不把外部營收結果當單方保證"],
        ],
    )

    doc.add_heading("九、老師必須先回答的12個問題", level=1)
    questions = [
        "半年後理想畫面是什麼：更多檢測、更多課程、蠟燭穩定回購，還是老師有更多時間？",
        "如果90天只能優先驗證一個產品與一條顧客旅程，選什麼？為什麼？",
        "母品牌正式是『希望之光』還是『希望之盒』？兩者是否有不同角色？",
        "大貴人9、小貴人9、大貴人7的正式價格、成本、頻道與差異是什麼？",
        "60分鐘諮詢正式價格已確認為1,980；創賦密碼、靈魂藍圖與掌運卡，是同一服務的三種工具，還是三個不同方案？",
        "21顆首購組的成本1,615包含什麼？是否含運、包材與諮詢？50組是否仍有效？",
        "蠟燭由誰製作、老師收到後實際完成哪些步驟？哪些畫面可拍、哪些說法可公開？",
        "十款蠟燭的成分、香味、燃燒時間、保存、消防、退換與敏感注意事項是什麼？",
        "腦意識觀察使用哪些儀器與資料？客戶會拿到什麼？不能回答或不適合服務的情況是什麼？",
        "老師能承接的回覆量、每週出鏡時間、檢測時段、課程名額與庫存產能是多少？",
        "上千份個案與800則訊息目前是否有取得行銷與公開故事同意？誰能合法、合宜地使用？",
        "品牌規劃者的固定費、績效歸因、獎金門檻、新產品分潤與終止條件如何書面化？",
    ]
    add_bullets(doc, questions)

    doc.add_heading("十、目前已完成與尚未完成", level=1)
    add_doc_table(
        doc,
        ["已完成初版", "仍需老師／團隊輸入"],
        [
            ["來源稽核、品牌架構、雙軌旅程、90天逐項路徑", "正式品牌名稱、品牌一句話與帳號決策"],
            ["23個SKU候選主檔、產品角色與安全版文案骨架", "正式價格、成本、規格、在售狀態與產品機制"],
            ["庫存、訂單、CRM、服務、回訪、內容、故事、LINE與分潤表格", "實物盤點、近3月交易、平台後台與名單資料"],
            ["40個任務化內容題目與10種LINE回覆情境", "真實案例、書面同意、拍攝素材與老師核准"],
            ["18項資料問題與來源採用／隔離紀錄", "逐項關閉待確認，不把建議當成已核准事實"],
        ],
    )
    p = doc.add_paragraph(style="Callout")
    p.add_run("這份初版的價值，不是替老師把答案猜完；而是把所有答案應該出現的位置、決策順序與後續營運路徑先建立起來。")

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.add_run("希望之光品牌規劃路徑圖｜2026-08-09｜[Codex] 初版｜待確認項目須由老師核准")
        footer_run.font.name = "Microsoft JhengHei"
        footer_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        footer_run.font.size = Pt(8)
        footer_run.font.color.rgb = RGBColor.from_string(MID_GREY)

    doc.core_properties.title = "希望之光品牌規劃路徑圖"
    doc.core_properties.author = "[Codex]"
    doc.core_properties.subject = "90天品牌整合、產品矩陣、顧客旅程與資料系統"
    doc.core_properties.comments = "依使用者提供來源整理；建議不等於品牌持有者核准。"
    doc.save(path)


def build_audit_markdown(path: Path) -> None:
    lines = [
        "# 希望之光｜資料盤點與待確認清單",
        "",
        "*Created: 2026-08-09 by [Codex]*",
        "",
        "本文件只記錄來源、衝突與下一步，不把未確認內容寫成既定事實。原始檔保持不動。",
        "",
        "## 優先處理順序",
        "",
        "1. 關閉品牌名稱、大小貴人價格與產品正式名稱；60分鐘諮詢價已確認為1,980。",
        "2. 現場盤點所有實體SKU，建立實際盤點日的期初異動。",
        "3. 補近三個月銷售與IG／LINE後台，建立真實基準後再談銷量KPI與分潤。",
        "4. 先用30-50份手寫個案試做數位化，確認同意、準確率與人工校正成本。",
        "5. 諧和機與所有健康相關文案先做安全改寫與必要審查。",
        "",
        "## 問題清單",
        "",
        "| ID | 嚴重度 | 類別 | 發現 | 建議 | 狀態 |",
        "|---|---|---|---|---|---|",
    ]
    for issue in DATA_ISSUES:
        safe = [str(value).replace("|", "／").replace("\n", " ") for value in issue]
        lines.append(f"| {safe[0]} | {safe[1]} | {safe[2]} | {safe[3]} | {safe[5]} | {safe[7]} |")
    lines += ["", "## 來源處理", "", "| 來源檔 | 用途 | 處理 | 說明 |", "|---|---|---|---|"]
    for source in SOURCE_AUDIT:
        safe = [str(value).replace("|", "／").replace("\n", " ") for value in source]
        lines.append(f"| {safe[0]} | {safe[1]} | {safe[2]} | {safe[3]} |")
    lines += [
        "",
        "## 隱私與資料治理",
        "",
        "- `希望之光客戶名單.xlsx` 的 MEGAWING 國外客戶分頁未匯入新CRM，因其公司抬頭與年份顯示可能屬舊業務資料。",
        "- 大小貴人經銷候選只匯入來源已有的最少資料；聯絡與行銷同意仍待補。",
        "- 個案故事必須分開記錄內部使用、匿名公開與具名公開同意，並保留撤回日期。",
        "",
    ]
    path.write_text("\n".join(lines), encoding="utf-8")


def validate_workbook(path: Path) -> dict[str, object]:
    wb = load_workbook(path, read_only=False, data_only=False)
    expected = [
        "00_使用說明", "01_品牌儀表板", "02_90天路徑", "03_產品主檔", "04_產品介紹",
        "05_庫存異動", "06_庫存總覽", "07_銷售訂單", "08_客戶CRM", "09_服務紀錄",
        "10_回訪任務", "11_內容主題庫", "12_內容成效", "13_希望故事庫", "14_LINE漏斗",
        "15_LINE回覆SOP", "16_分潤試算", "17_品牌決策", "18_資料核對", "19_選單",
    ]
    missing = [name for name in expected if name not in wb.sheetnames]
    order_ok = wb.sheetnames == expected
    formula_count = 0
    error_formulas: list[str] = []
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if isinstance(cell.value, str) and cell.value.startswith("="):
                    formula_count += 1
                    if "#REF!" in cell.value.upper():
                        error_formulas.append(f"{ws.title}!{cell.coordinate}")
    hidden_ok = wb["19_選單"].sheet_state == "hidden"
    result = {
        "sheets": len(wb.sheetnames),
        "missing_sheets": missing,
        "sheet_order_ok": order_ok,
        "formula_cells": formula_count,
        "ref_error_formulas": error_formulas,
        "hidden_lists_sheet": hidden_ok,
        "size_bytes": path.stat().st_size,
    }
    wb.close()
    return result


def main() -> None:
    OUTPUT.mkdir(parents=True, exist_ok=True)
    xlsx_path = OUTPUT / "希望之光_品牌營運管理系統_202608.xlsx"
    docx_path = OUTPUT / "希望之光_品牌規劃路徑圖_202608.docx"
    audit_path = OUTPUT / "資料盤點與待確認清單.md"
    validation_dir = HERE / "validation"
    validation_dir.mkdir(parents=True, exist_ok=True)
    report_path = validation_dir / "build_validation.json"

    build_workbook(xlsx_path)
    from build_simple_tools import build_solo_roadmap

    build_solo_roadmap(docx_path)
    build_audit_markdown(audit_path)
    report = {
        "created": date.today().isoformat(),
        "workbook": validate_workbook(xlsx_path),
        "docx_size_bytes": docx_path.stat().st_size,
        "audit_size_bytes": audit_path.stat().st_size,
    }
    report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
